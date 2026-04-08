# ==============================================
# phase4_cover_letter.py
# Staged cover letter generator for a specific role.
# Mirrors the Phase 4 resume staged file workflow.
#
# Stage 1: Generate cl_stage1_draft.txt (both sections)
# Stage 4: Convert cl_stage4_final.txt to .docx
#
# Staged file flow (data/job_packages/[role]/):
#   cl_stage1_draft.txt     <- script writes (--stage 1)
#   cl_stage2_approved.txt  <- user edits and saves manually
#   cl_stage3_review.txt    <- check_cover_letter.py writes
#   cl_stage4_final.txt     <- user saves manually (source for docx)
#
# Docx output: resumes/tailored/[role]/[role]_CoverLetter.docx
#
# Usage:
#   python scripts/phase4_cover_letter.py --stage 1 --role BAH_LCI_MBSE
#   python scripts/phase4_cover_letter.py --stage 4 --role BAH_LCI_MBSE
# ==============================================

import os
import sys
import re
import argparse
from datetime import datetime
from anthropic import Anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches

from scripts.utils.pii_filter import strip_pii

load_dotenv()

# ==============================================
# CONFIGURATION
# ==============================================

JOBS_PACKAGES_DIR = "data/job_packages"
RESUMES_TAILORED_DIR = "resumes/tailored"
CANDIDATE_BACKGROUND_PATH = "context/CANDIDATE_BACKGROUND.md"
RESUME_TEMPLATE = "templates/resume_template.docx"
MODEL = "claude-sonnet-4-20250514"

CL_STAGE1 = "cl_stage1_draft.txt"
CL_STAGE4 = "cl_stage4_final.txt"

# ==============================================
# HELPER FUNCTIONS
# ==============================================

def check_overwrite(filepath, label):
    """Warn if output file already exists."""
    if os.path.exists(filepath):
        print(f"\nWARNING: {label} already exists: {filepath}")
        response = input("  Overwrite? (y/n): ").strip().lower()
        if response != 'y':
            print("  Cancelled. Existing file preserved.")
            return False
    return True


def load_resume_bullets(stage4_path, stage2_path):
    """
    Extract resume content from stage4_final.txt (or stage2_approved.txt).
    Stage files are source of truth \u2013 .docx is presentation only.
    Returns (clean_text str, source filename).
    """
    path = source = None
    if os.path.exists(stage4_path):
        path, source = stage4_path, "stage4_final.txt"
    elif os.path.exists(stage2_path):
        path, source = stage2_path, "stage2_approved.txt"
    else:
        return None, None

    with open(path, encoding='utf-8') as f:
        content = f.read()

    clean_lines = []
    for line in content.split('\n'):
        stripped = line.strip()
        if (stripped.startswith('=') or stripped.startswith('STAGE') or
                stripped.startswith('Role:') or stripped.startswith('Generated:') or
                stripped.startswith('INSTRUCTIONS') or stripped.startswith('Save as') or
                stripped.startswith('END OF')):
            continue
        cleaned = re.sub(r'\s*\[Source:[^\]]*\]', '', stripped)
        cleaned = re.sub(r'\s*\[Theme:[^\]]*\]', '', cleaned)
        cleaned = re.sub(r'\s*\[VERIFY[^\]]*\]', '', cleaned)
        cleaned = re.sub(r'\s*\[FLAGGED[^\]]*\]', '', cleaned)
        if cleaned:
            clean_lines.append(cleaned)

    return '\n'.join(clean_lines), source


def extract_coverage_gaps(stage3_review_path):
    """
    Parse COVERAGE GAPS section from resume stage3_review.txt.
    Returns list of gap topic strings; empty list if file absent or section missing.
    """
    if not os.path.exists(stage3_review_path):
        return []

    with open(stage3_review_path, encoding='utf-8') as f:
        content = f.read()

    gaps = []
    in_gaps = False
    for line in content.split('\n'):
        stripped = line.strip()
        if stripped.startswith('COVERAGE GAPS:'):
            in_gaps = True
            continue
        if in_gaps:
            if not stripped:
                break  # blank line ends the gaps section
            if stripped.startswith('- '):
                gap_text = stripped[2:].strip().replace('**', '')
                if ':' in gap_text:
                    gap_text = gap_text.split(':')[0].strip()
                if gap_text:
                    gaps.append(gap_text)
            elif not stripped.startswith('-'):
                break  # new section header

    return gaps


def extract_hiring_manager(client, jd):
    """
    Try to find hiring manager name in JD via Claude.
    Non-interactive \u2013 returns name string or 'Hiring Manager'.
    """
    try:
        response = client.messages.create(
            model=MODEL,
            max_tokens=50,
            messages=[{
                "role": "user",
                "content": (
                    "Extract the hiring manager's name from this job description "
                    "if one is explicitly stated. Reply with ONLY the name, or "
                    "reply with exactly 'NOT FOUND' if no name is present.\n\n"
                    f"{jd[:3000]}"
                )
            }]
        )
        result = response.content[0].text.strip()
        if result and result.upper() != "NOT FOUND":
            return result
    except Exception:
        pass
    return "Hiring Manager"


def fix_dashes(text):
    """Replace em dashes with en dashes throughout."""
    return text.replace('\u2014', '\u2013')

# ==============================================
# SYSTEM PROMPT
# ==============================================

SYSTEM_PROMPT = """You are an expert career consultant specializing in defense and
aerospace systems engineering. You help senior engineers write compelling,
authentic cover letters.

You always:
- Ground all claims strictly in the candidate's confirmed experience provided
- Use en dashes (\u2013), never em dashes (\u2014)
- Write in a direct, confident, professional tone
- Tailor content to the specific role and company from the JD

You never:
- Invent metrics, outcomes, or experience not in the provided materials
- Fabricate unverifiable achievements
- Use filler opener phrases ("I am excited to apply...", "I am writing to express...")
- Use em dashes (\u2014) \u2013 use en dashes (\u2013) only
- Reference CompTIA Security+ (the candidate's certification has lapsed)
- Use "Active TS/SCI" \u2013 always "Current TS/SCI" when the candidate is between employers"""

# ==============================================
# API CALLS \u2013 STAGE 1
# ==============================================

def generate_traditional_letter(client, jd, bullets, background,
                                  hiring_manager, gaps):
    """Call 1 \u2013 Generate 3-4 paragraph traditional cover letter."""
    now = datetime.now()
    today = f"{now.strftime('%B')} {now.day}, {now.year}"

    gaps_block = ""
    if gaps:
        gaps_block = (
            "\nCONFIRMED GAPS FROM RESUME REVIEW \u2013 do NOT claim or imply these:\n"
            + "\n".join(f"  \u2013 {g}" for g in gaps)
            + "\n"
        )

    bullets_text = bullets[:3000] if bullets else "Not available \u2013 use candidate background only."
    prompt = f"""Write a traditional cover letter for this job application.

HIRING MANAGER: {hiring_manager}
TODAY'S DATE: {today}

JOB DESCRIPTION:
{jd}

CANDIDATE BACKGROUND (PII removed):
{background[:2000]}

RESUME BULLETS FOR THIS ROLE:
{bullets_text}
{gaps_block}
RULES:
- Full business letter format: date line, company name as inside address,
  salutation, 3-4 body paragraphs, professional closing
- Opening paragraph: state the specific role and lead immediately with the
  strongest relevant credential \u2013 no generic opener sentences
- Middle paragraphs: connect specific resume bullets to key JD requirements
- Closing paragraph: concise call to action, one sentence
- Salutation: "Dear {hiring_manager},"
- En dashes (\u2013) only, never em dashes (\u2014)
- No fabricated metrics or experience not in the provided materials
- No CompTIA Security+ reference (certification lapsed)
- Use "Current TS/SCI" \u2013 never "Active TS/SCI"
- Write in first person

Write the letter starting with the date line ({today})."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=1500,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text


def generate_application_paragraph(client, jd, bullets, background, gaps):
    """Call 2 \u2013 Generate 150-250 word application field paragraph."""
    gaps_block = ""
    if gaps:
        gaps_block = (
            "\nCONFIRMED GAPS \u2013 do NOT claim or imply these:\n"
            + "\n".join(f"  \u2013 {g}" for g in gaps)
            + "\n"
        )

    prompt = f"""Write a short application paragraph for this role.

JOB DESCRIPTION:
{jd}

CANDIDATE BACKGROUND (PII removed):
{background[:2000]}

RESUME BULLETS FOR THIS ROLE:
{bullets[:3000] if bullets else 'Not available \u2013 use candidate background only.'}
{gaps_block}
INSTRUCTIONS:
- Determine whether a condensed cover letter or an elevator pitch / interest
  statement better fits this role. Use the JD tone and company type as your
  guide: startups want an interest statement; large defense primes and
  government contractors want a condensed letter.
- Write 150-250 words suitable for pasting into a plain-text application field.
- No letterhead, no date, no salutation, no closing signature.
- Prose paragraphs only.
- En dashes (\u2013) only, never em dashes (\u2014).
- No fabricated metrics or experience.
- No CompTIA Security+ reference.
- Use "Current TS/SCI" if clearance is mentioned \u2013 never "Active TS/SCI."

Write the paragraph now."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=600,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.content[0].text

# ==============================================
# OUTPUT ASSEMBLY \u2013 STAGE 1
# ==============================================

def build_stage1_txt(role, section1, section2, resume_source):
    """Assemble cl_stage1_draft.txt content with section delimiters."""
    timestamp = datetime.now().strftime('%d %b %Y %H:%M')
    lines = [
        "=" * 60,
        "COVER LETTER DRAFT",
        "=" * 60,
        f"Role: {role}",
        f"Generated: {timestamp}",
        f"Resume source: {resume_source if resume_source else 'Not found'}",
        "Note: PII stripped from all API calls.",
        "INSTRUCTIONS:",
        f"  1. Review both sections \u2013 verify all claims against your experience",
        f"  2. Make edits and save as cl_stage2_approved.txt",
        f"  3. Run: python scripts/check_cover_letter.py --role {role}",
        f"  4. Review cl_stage3_review.txt, make final edits, save as cl_stage4_final.txt",
        f"  5. Run: python scripts/phase4_cover_letter.py --stage 4 --role {role}",
        "=" * 60,
        "",
        "## COVER LETTER",
        "",
        section1,
        "",
        "## APPLICATION PARAGRAPH",
        "",
        section2,
        "",
        "=" * 60,
        "END OF COVER LETTER DRAFT",
        "=" * 60,
    ]
    return "\n".join(lines)

# ==============================================
# STAGE 4 \u2013 PARSE AND BUILD DOCX
# ==============================================

def parse_stage4(content):
    """
    Parse cl_stage4_final.txt into (section1, section2) strings.
    Splits on ## COVER LETTER and ## APPLICATION PARAGRAPH markers.
    Raises ValueError if markers are missing.
    """
    if '## COVER LETTER' not in content:
        raise ValueError(
            "cl_stage4_final.txt is missing the '## COVER LETTER' marker.\n"
            "  Ensure the file was derived from cl_stage1_draft.txt and "
            "the section markers were preserved."
        )
    if '## APPLICATION PARAGRAPH' not in content:
        raise ValueError(
            "cl_stage4_final.txt is missing the '## APPLICATION PARAGRAPH' marker.\n"
            "  Ensure the file was derived from cl_stage1_draft.txt and "
            "the section markers were preserved."
        )

    after_cl = content.split('## COVER LETTER', 1)[1]
    section1_raw, section2_raw = after_cl.split('## APPLICATION PARAGRAPH', 1)

    # Strip trailing footer from section2
    section2 = section2_raw.split('=' * 60)[0].strip()
    section1 = section1_raw.strip()

    return section1, section2


def build_cover_letter_docx(output_path, section1, section2):
    """
    Generate cover letter .docx from parsed stage4 content.
    Page 1: traditional letter. Page break. Page 2: application paragraph.
    No labels or metadata in docx. US Letter page size.
    """
    # Load template for font styles; clear pre-filled body content
    if os.path.exists(RESUME_TEMPLATE):
        doc = Document(RESUME_TEMPLATE)
        print(f"  Using template: {RESUME_TEMPLATE}")
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag != 'sectPr':  # Keep section properties
                body.remove(child)
    else:
        print(f"  WARNING: Template not found at {RESUME_TEMPLATE} \u2013 using basic formatting")
        doc = Document()

    # Page setup: US Letter, 1" margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    def add_normal(text):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)

    def add_spacer():
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def render_section(text):
        """Render section text as Normal paragraphs, preserving blank lines."""
        for line in text.split('\n'):
            stripped = line.strip()
            if stripped:
                add_normal(stripped)
            else:
                add_spacer()

    # Page 1: Traditional cover letter
    render_section(section1)

    # Page break between sections
    doc.add_page_break()

    # Page 2: Application paragraph
    render_section(section2)

    doc.save(output_path)

# ==============================================
# EXTRACTED STAGE FUNCTIONS (testable)
# ==============================================

def run_cl_stage1(client, jd_text, resume_text, background_text, output_path):
    """
    Generate cl_stage1_draft.txt content from inputs and write to output_path.

    Strips PII from resume_text and jd_text before any API call.
    Makes two API calls: traditional letter + application paragraph.
    Writes assembled draft to output_path.
    """
    jd_clean = strip_pii(jd_text)
    bullets_clean = strip_pii(resume_text)
    background_clean = strip_pii(background_text)

    hiring_manager = extract_hiring_manager(client, jd_clean)

    section1 = generate_traditional_letter(
        client, jd_clean, bullets_clean, background_clean, hiring_manager, []
    )
    section2 = generate_application_paragraph(
        client, jd_clean, bullets_clean, background_clean, []
    )

    section1 = fix_dashes(section1)
    section2 = fix_dashes(section2)

    output_text = build_stage1_txt("test_role", section1, section2, "stage2_approved.txt")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(output_text)


def run_cl_stage4(cl_text, output_path):
    """
    Generate a cover letter .docx from plain text (no API call).

    Writes each line of cl_text as a Normal paragraph to output_path.
    """
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    for line in cl_text.split('\n'):
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.add_run(stripped)
        else:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

    doc.save(output_path)

# ==============================================
# MAIN
# ==============================================

def main():
    parser = argparse.ArgumentParser(description='Phase 4 Cover Letter Generator')
    parser.add_argument('--stage', type=int, required=True, choices=[1, 4],
                        help='Stage to run (1=generate draft, 4=build docx)')
    parser.add_argument('--role', type=str, required=True,
                        help='Role package folder name (e.g. BAH_LCI_MBSE)')
    args = parser.parse_args()

    STAGE = args.stage
    ROLE = args.role

    PACKAGE_DIR = os.path.join(JOBS_PACKAGES_DIR, ROLE)
    JD_PATH = os.path.join(PACKAGE_DIR, "job_description.txt")
    STAGE3_REVIEW_PATH = os.path.join(PACKAGE_DIR, "stage3_review.txt")   # resume review
    STAGE4_RESUME_PATH = os.path.join(PACKAGE_DIR, "stage4_final.txt")    # resume bullets
    STAGE2_RESUME_PATH = os.path.join(PACKAGE_DIR, "stage2_approved.txt") # fallback bullets
    CL_STAGE1_PATH = os.path.join(PACKAGE_DIR, CL_STAGE1)
    CL_STAGE4_PATH = os.path.join(PACKAGE_DIR, CL_STAGE4)
    TAILORED_DIR = os.path.join(RESUMES_TAILORED_DIR, ROLE)
    DOCX_PATH = os.path.join(TAILORED_DIR, f"{ROLE}_CoverLetter.docx")

    if STAGE == 1:
        print("=" * 60)
        print("PHASE 4 \u2013 COVER LETTER GENERATOR \u2013 STAGE 1")
        print("=" * 60)
        print(f"Role: {ROLE}")
        print(f"Package: {PACKAGE_DIR}")

        errors = []
        if not os.path.exists(PACKAGE_DIR):
            errors.append(f"Job package folder not found: {PACKAGE_DIR}")
        if not os.path.exists(JD_PATH):
            errors.append(f"job_description.txt not found in {PACKAGE_DIR}")
        if not os.path.exists(CANDIDATE_BACKGROUND_PATH):
            errors.append(f"CANDIDATE_BACKGROUND.md not found: {CANDIDATE_BACKGROUND_PATH}")

        if errors:
            print(f"\nERRORS \u2013 cannot proceed:")
            for e in errors:
                print(f"  {e}")
            sys.exit(1)

        if not check_overwrite(CL_STAGE1_PATH, CL_STAGE1):
            sys.exit(0)

        print("\nLoading data...")

        with open(JD_PATH, encoding='utf-8') as f:
            jd = f.read()
        print(f"  JD loaded: {len(jd)} chars")

        with open(CANDIDATE_BACKGROUND_PATH, encoding='utf-8') as f:
            raw_background = f.read()
        background = strip_pii(raw_background)
        print("  Candidate background loaded and PII stripped.")

        resume_bullets, resume_source = load_resume_bullets(
            STAGE4_RESUME_PATH, STAGE2_RESUME_PATH
        )
        if resume_bullets:
            print(f"  Resume loaded from {resume_source}")
        else:
            print("  WARNING: No resume stage file found \u2013 cover letter will use background only")
            resume_source = None

        gaps = extract_coverage_gaps(STAGE3_REVIEW_PATH)
        if gaps:
            print(f"  Coverage gaps loaded: {len(gaps)} gaps (from stage3_review.txt)")
            for g in gaps:
                print(f"    \u2013 {g}")
        else:
            print("  No stage3_review.txt found \u2013 gap filtering skipped")

        client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

        print("\n  Checking JD for hiring manager name...")
        hiring_manager = extract_hiring_manager(client, jd)
        if hiring_manager == "Hiring Manager":
            print("  Hiring manager: Hiring Manager (default \u2013 not found in JD)")
        else:
            print(f"  Hiring manager: {hiring_manager}")

        bullets_for_api = strip_pii(resume_bullets) if resume_bullets else ""

        print("\nGenerating cover letter...")
        print("  [1/2] Traditional cover letter...")
        section1 = generate_traditional_letter(
            client, jd, bullets_for_api, background, hiring_manager, gaps
        )

        print("  [2/2] Application paragraph...")
        section2 = generate_application_paragraph(
            client, jd, bullets_for_api, background, gaps
        )

        # Enforce en dashes
        section1 = fix_dashes(section1)
        section2 = fix_dashes(section2)

        print(f"\nWriting {CL_STAGE1}...")
        output_text = build_stage1_txt(ROLE, section1, section2, resume_source)
        with open(CL_STAGE1_PATH, 'w', encoding='utf-8') as f:
            f.write(output_text)

        print(f"\n{'=' * 60}")
        print("STAGE 1 COMPLETE")
        print(f"{'=' * 60}")
        print(f"Draft: {CL_STAGE1_PATH}")
        print(f"\nNext steps:")
        print(f"  1. Open {CL_STAGE1} \u2013 verify all claims are grounded")
        print(f"  2. Make edits and save as cl_stage2_approved.txt")
        print(f"  3. Run: python scripts/check_cover_letter.py --role {ROLE}")
        print(f"  4. Review cl_stage3_review.txt, save final as cl_stage4_final.txt")
        print(f"  5. Run: python scripts/phase4_cover_letter.py --stage 4 --role {ROLE}")
        print(f"{'=' * 60}")

    elif STAGE == 4:
        print("=" * 60)
        print("PHASE 4 \u2013 COVER LETTER GENERATOR \u2013 STAGE 4")
        print("=" * 60)
        print(f"Role: {ROLE}")
        print(f"Source: {CL_STAGE4_PATH}")
        print(f"Output: {DOCX_PATH}")

        errors = []
        if not os.path.exists(PACKAGE_DIR):
            errors.append(f"Job package folder not found: {PACKAGE_DIR}")
        if not os.path.exists(CL_STAGE4_PATH):
            errors.append(
                f"{CL_STAGE4} not found in {PACKAGE_DIR}\n"
                f"  Complete stages 1\u20133 first, then save final version as {CL_STAGE4}."
            )

        if errors:
            print(f"\nERRORS \u2013 cannot proceed:")
            for e in errors:
                print(f"  {e}")
            sys.exit(1)

        if not check_overwrite(DOCX_PATH, f"{ROLE}_CoverLetter.docx"):
            sys.exit(0)

        # Ensure tailored output directory exists
        os.makedirs(TAILORED_DIR, exist_ok=True)

        print(f"\nParsing {CL_STAGE4}...")
        with open(CL_STAGE4_PATH, encoding='utf-8') as f:
            stage4_content = f.read()

        try:
            section1, section2 = parse_stage4(stage4_content)
        except ValueError as e:
            print(f"\nERROR: {e}")
            sys.exit(1)

        print(f"  Cover letter: {len(section1)} chars")
        print(f"  Application paragraph: {len(section2)} chars")

        print(f"\nBuilding {ROLE}_CoverLetter.docx...")
        try:
            build_cover_letter_docx(DOCX_PATH, section1, section2)
        except Exception as e:
            print(f"\nERROR: Docx generation failed: {e}")
            sys.exit(1)

        print(f"\n{'=' * 60}")
        print("STAGE 4 COMPLETE")
        print(f"{'=' * 60}")
        print(f"Docx: {DOCX_PATH}")
        print(f"\nNext steps:")
        print(f"  1. Open in Word \u2013 verify formatting and no em dashes")
        print(f"  2. Confirm page 1 is the letter, page 2 is the application paragraph")
        print(f"  3. Update hiring manager name/address if placeholder text remains")
        print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
