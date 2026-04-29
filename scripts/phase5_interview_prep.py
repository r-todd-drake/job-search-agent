# ==============================================
# phase5_interview_prep.py
# Generates stage-aware interview preparation
# materials for a specific job application.
#
# Improvements over v2:
#   - --interview_stage parameter drives all
#     section content (recruiter / hiring_manager
#     / team_panel)
#   - "Introduce Yourself" section tailored per
#     stage from candidate_profile.md
#   - Short tenure explanation prepended to
#     gap prep section at all stages
#   - Stage-specific questions prompt per audience
#   - --dry_run flag for profile validation
#   - Stage-specific output filenames
#
# Outputs to data/job_packages/[role]/:
#   interview_prep_[stage].txt
#   interview_prep_[stage].docx
#
# Sections:
#   1:   Company & Role Brief (web-informed)
#   1.5: Introduce Yourself (stage-tailored)
#   2:   Story Bank (library-grounded)
#   3:   Gap Preparation (stage-conditional)
#   4:   Questions to Ask (stage-specific)
#
# Usage:
#   python -m scripts.phase5_interview_prep \
#     --role Acme_SE_Systems \
#     --interview_stage hiring_manager
#   python -m scripts.phase5_interview_prep \
#     --role Acme_SE_Systems --dry_run
# ==============================================

import os
import sys
import re
import json
import argparse
from datetime import datetime
from anthropic import Anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add project root to path for utils import
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from scripts.utils.pii_filter import strip_pii
import scripts.interview_library_parser as _ilp
from scripts.config import (
    JOBS_PACKAGES_DIR,
    CANDIDATE_PROFILE_PATH,
    EXPERIENCE_LIBRARY_JSON as EXPERIENCE_LIBRARY,
    MODEL_SONNET as MODEL,
)

load_dotenv()

# ==============================================
# CONFIGURATION
# ==============================================

OUTPUT_FILENAME = "interview_prep.txt"
OUTPUT_DOCX_FILENAME = "interview_prep.docx"

# ==============================================
# SYSTEM PROMPT
# ==============================================

SYSTEM_PROMPT = """You are an expert career coach specializing in defense and aerospace
systems engineering. You help senior engineers prepare for technical interviews.

You always:
- Ground all stories and claims in the candidate's confirmed experience only
- Use employer attribution in stories ("During my time at G2 OPS..." or
  "When I was supporting Shield AI...")
- Frame stories using STAR format with factually accurate details
- Give honest gap assessments -- never suggest claiming experience not held
- Keep language professional, direct, and confident
- Use en dashes, never em dashes

You never:
- Invent metrics, outcomes, or experience not in the provided background
- Suggest the candidate overstate their role or involvement
- Use vague or generic advice"""

# ==============================================
# STAGE PROFILE CONSTANTS
# ==============================================

_QUESTIONS_RECRUITER = """
Generate exactly 4 questions for a candidate to ask at the end of a recruiter screen.
The questions should signal that the candidate is serious, prepared, and professional
without asking technical or program-specific questions the recruiter cannot answer.

Question categories to cover:
1. Company direction, growth areas, or recent news the candidate can reference naturally
2. Culture, team environment, or what makes people stay at the company
3. Interview process -- who is next in the process, what will they be evaluating, timeline to decision
4. One logistics or role clarification question if anything material remains unaddressed (clearance, location, remote/onsite)

Constraints:
- Questions must be answerable by a recruiter without program or technical knowledge
- Do not ask about architecture, tooling, program pain points, or technical environment
- Do not ask questions already answered in the job description
- Tone: engaged, collegial, unhurried -- not transactional
- Format: numbered list, each question followed by one sentence explaining what it signals to the interviewer
"""


_QUESTIONS_HIRING_MANAGER = """
Generate exactly 4 questions for a candidate to ask at the end of a hiring manager interview
for a defense systems engineering role. The questions should signal program awareness,
results orientation, and genuine interest in the manager's vision.

Question categories to cover:
1. Current program pain points -- where is the pressure coming from (schedule, architecture debt, stakeholder friction)?
2. What the team needs that it does not currently have -- what gap does this hire fill?
3. What success looks like at 6 months versus what disappointment looks like -- what are the real expectations?
4. The hiring manager's vision for where the technical effort (MBSE, architecture, or the relevant discipline) goes from here

Constraints:
- Questions must require insider knowledge to answer well -- not answerable from the job description alone
- Do not ask about company culture, interview process, or logistics -- those belong in the recruiter screen
- Do not ask questions that make the candidate sound uncertain about fit or qualifications
- Tone: peer-level engagement with someone senior -- collaborative, direct, curious about the problem
- Format: numbered list, each question followed by one sentence explaining what it signals to the interviewer
"""


_QUESTIONS_TEAM_PANEL = """
Generate exactly 4 questions for a candidate to ask at the end of a team panel interview
for a defense systems engineering role. The questions should signal technical credibility,
process fluency, and peer-level awareness of where the hard work actually lives.

Question categories to cover:
1. Day-to-day working environment -- tools cadence, model governance, or workflow specifics
   that only someone who has done this work before would think to ask about
2. Where the hard interface or integration problems are concentrated right now
3. What processes are working well and what is still being figured out -- invites honest answer
4. How the team handles disagreements on architecture or design decisions -- signals maturity
   and interest in team dynamics at a working level

Constraints:
- Questions must require hands-on program knowledge to answer -- not answerable from the JD alone
- Do not ask about company direction, culture, salary, or interview process
- Do not ask questions that sound like a candidate evaluating risk -- ask like a peer evaluating the work
- Avoid questions that could be read as critical of the program or the team
- Tone: direct, collegial, technically confident -- peer to peer, not candidate to evaluator
- Format: numbered list, each question followed by one sentence explaining what it signals to the interviewer
"""


_PEER_FRAME_INSTRUCTIONS = """
For the gap identified above, generate a Peer Frame response suitable for delivery
to a working-level engineer in a team panel interview.

The Peer Frame must:
1. Acknowledge the specific gap honestly -- no softening, no hedging
2. Demonstrate that the candidate understands why this gap matters operationally,
   not just that the gap exists -- show awareness of where the friction point actually lives
3. Pivot to a question or observation that signals domain fluency -- the candidate
   should sound like someone who has worked adjacent to this problem before

Tone: direct and collegial -- peer to peer, not candidate to evaluator.
Do not use polished redirect language or reassurance framing -- those belong in
the hiring manager response, not here.
A Peer Frame that ends with a genuine question is strongly preferred over one
that ends with a reassurance statement.

Length: 2-3 sentences maximum.
Label the output: Peer Frame:
"""


STAGE_PROFILES = {
    "recruiter": {
        "label": "Recruiter Screen",
        "description": "Short screen \u2013 confirm fit, do not volunteer gaps or technical depth.",
        "story_count": "1-2",
        "story_depth": "headline",
        "gap_behavior": "omit",
        "salary_in_section1": False,
        "section1_focus": "recruiter",
        "questions_prompt": _QUESTIONS_RECRUITER,
    },
    "hiring_manager": {
        "label": "Hiring Manager Interview",
        "description": "60+ min interview \u2013 lead with program context awareness and collaborative framing.",
        "story_count": "3-4",
        "story_depth": "full",
        "gap_behavior": "note",
        "salary_in_section1": True,
        "section1_focus": "hiring_manager",
        "questions_prompt": _QUESTIONS_HIRING_MANAGER,
    },
    "team_panel": {
        "label": "Team Panel Interview",
        "description": "90 min to 3 hr group interview \u2013 lead with technical specificity and process fluency.",
        "story_count": "4-6",
        "story_depth": "full_technical",
        "gap_behavior": "full_peer",
        "salary_in_section1": False,
        "section1_focus": "team_panel",
        "questions_prompt": _QUESTIONS_TEAM_PANEL,
        "peer_frame_prompt": _PEER_FRAME_INSTRUCTIONS,
    },
}

VALID_STAGES = list(STAGE_PROFILES.keys())

# ==============================================
# OUTPUT PATH HELPER
# ==============================================

def _output_paths(package_dir, stage):
    """Return (txt_path, docx_path) for the given stage."""
    return (
        os.path.join(package_dir, f"interview_prep_{stage}.txt"),
        os.path.join(package_dir, f"interview_prep_{stage}.docx"),
    )


def extract_profile_section(profile_text, header):
    """
    Extract a ## HEADER section from candidate_profile.md text.
    Returns the section body (stripped), or empty string if header not found.
    """
    marker = f"## {header}"
    if marker not in profile_text:
        return ""
    start = profile_text.find(marker) + len(marker)
    next_header = profile_text.find("\n## ", start)
    end = next_header if next_header > 0 else len(profile_text)
    return profile_text[start:end].strip()


def _build_section1_prompt(jd, salary_data, profile, salary_actuals=None):
    """Build the Section 1 company brief prompt, parameterized by stage profile."""
    _stage_instructions = {
        "recruiter": (
            "Focus on:\n"
            "- Company overview (3-4 sentences): what they do, defense/government focus, scale\n"
            "- Culture signals: what employees say about the environment and retention\n"
            "- Recent news: contracts, programs, or announcements relevant to this role\n"
            "- Interview process context: who typically interviews next, what they evaluate\n"
            "Omit salary guidance. Do not include detailed program or technical content."
        ),
        "hiring_manager": (
            "Focus on:\n"
            "- Full company overview (3-4 sentences): mission, defense/government business, scale\n"
            "- Business unit deep-dive (2-3 sentences): specific unit, programs, stakeholders\n"
            "- Program pain points: based on JD language, what problems is this role solving?\n"
            "- Role in context: day-to-day responsibilities inferred from JD\n"
            "Include salary guidance block."
        ),
        "team_panel": (
            "Focus on:\n"
            "- Company overview: CONDENSED to 2-3 sentences only -- panel members know the company\n"
            "- Program-specific context: mission area, technical environment, active programs from JD\n"
            "- Technical environment: tools, methodologies, and stack signals in JD language\n"
            "Omit salary guidance. Omit general culture content."
        ),
    }

    salary_block = ""
    if profile["salary_in_section1"]:
        if salary_actuals:
            min_val = salary_actuals.get("range_given_min")
            max_val = salary_actuals.get("range_given_max")
            anchor = salary_actuals.get("candidate_anchor")
            floor_ = salary_actuals.get("candidate_floor")
            act_stage = salary_actuals.get("stage", "prior stage")
            act_date = salary_actuals.get("interview_date", "prior interview")
            min_str = f"${min_val:,.0f}" if min_val else "not recorded"
            max_str = f"${max_val:,.0f}" if max_val else "not recorded"
            anchor_str = f"${anchor:,.0f}" if anchor else "not recorded"
            floor_str = f"${floor_:,.0f}" if floor_ else "not recorded"
            salary_block = (
                f"\nSALARY ACTUALS (reported from {act_stage} on {act_date} -- use these, not estimates):\n"
                f"Range given by interviewer: {min_str} -- {max_str}\n"
                f"Candidate anchor stated: {anchor_str}\n"
                f"Candidate floor: {floor_str}\n"
                f"Note: these are reported actuals from a prior interview for this role. "
                f"Present them as confirmed data, not as analysis.\n"
                f"IMPORTANT: Do not re-derive or add salary guidance from the job description. "
                f"Use ONLY the salary actuals block above as written.\n"
            )
        else:
            salary_block = (
                f"\nSALARY & LEVEL CONTEXT:\n"
                f"JD posted range: {salary_data['text'] if salary_data['found'] else 'Not found in JD'}\n"
                f"[1-2 sentences on what level this represents and where initial offers land.]\n\n"
                f"SALARY EXPECTATIONS GUIDANCE:\n"
                f"{salary_data['guidance'] if salary_data['found'] else 'Research market rate before interview.'}\n"
            )

    return (
        f"Research this company and role, then generate an interview prep brief "
        f"for a {profile['label']}.\n\n"
        f"JOB DESCRIPTION:\n{jd[:2500]}\n\n"
        f"Use the web_search tool to find current information about this company.\n\n"
        f"Stage-specific instructions:\n{_stage_instructions[profile['section1_focus']]}\n"
        f"{salary_block}\n"
        f"Format your brief with ALL-CAPS section headers followed by a colon "
        f"(e.g., 'COMPANY OVERVIEW:'). Include only sections relevant to this stage."
    )


def _build_intro_prompt(intro_monologue, profile):
    """Build the 'Introduce Yourself' tailoring prompt, parameterized by stage profile."""
    _tailoring = {
        "recruiter": (
            "2-3 sentences, high-level",
            "overall fit and interest in the role -- confirm you are not a risk",
        ),
        "hiring_manager": (
            "3-4 sentences, program-context aware",
            "program experience and collaborative working style",
        ),
        "team_panel": (
            "4-5 sentences, technically grounded",
            "specific tools, methodologies, and day-to-day peer-relevant experience",
        ),
    }
    length_guidance, emphasis = _tailoring[profile["section1_focus"]]

    return (
        f"The candidate has a prepared introduction for 'Tell me about yourself.'\n\n"
        f"BASE INTRODUCTION:\n{intro_monologue}\n\n"
        f"Tailor this introduction for a {profile['label']} interview.\n"
        f"- Length: {length_guidance}\n"
        f"- Emphasis: {emphasis}\n"
        f"- Register: appropriate for this audience ({profile['description']})\n\n"
        f"Rules:\n"
        f"- Keep all factual content present in the base text\n"
        f"- Do not add experience, credentials, or claims not in the base text\n"
        f"- Return the tailored introduction as flowing prose (1-2 short paragraphs max)\n"
        f"- Do not add headers or labels -- return only the introduction text itself"
    )

def _build_section2_prompt(jd, story_context, candidate_profile, profile, library_seeds=None):
    """Build the Section 2 story bank prompt, parameterized by stage profile."""
    _depth_instructions = {
        "headline": (
            "STORY DEPTH: Headline only.\n"
            "- Provide story headline + one-sentence result for each story\n"
            "- Do NOT expand to full STAR format\n"
            "- Omit 'If probed' branch"
        ),
        "full": (
            "STORY DEPTH: Full STAR with probe branch.\n"
            "- Full Situation / Task / Action / Result for each story\n"
            "- Include one 'If probed' branch per story (one additional sentence)"
        ),
        "full_technical": (
            "STORY DEPTH: Full STAR with technical specificity.\n"
            "- Full Situation / Task / Action / Result for each story\n"
            "- Use tool-specific language (name tools, models, frameworks used)\n"
            "- Include peer-credible detail a working engineer would recognize\n"
            "- Include one 'If probed' branch per story"
        ),
    }

    _gap_instructions = {
        "omit": "GAP FRAMING: Do NOT reference gaps or limitations in any story framing.",
        "note": (
            "GAP FRAMING: Where a story might brush against a known gap, "
            "include a brief one-sentence awareness note."
        ),
        "full": "GAP FRAMING: Integrate full gap awareness into story framing where relevant.",
        "full_peer": (
            "GAP FRAMING: Integrate full gap awareness into story framing where relevant, "
            "with peer-level directness."
        ),
    }

    role_fit_instruction = (
        "2 sentences only -- lead with strongest fit signal."
        if profile["story_depth"] == "headline"
        else "3-4 honest sentences -- genuine strengths and real gaps."
    )

    seed_block = ""
    if library_seeds:
        seed_lines = [
            "\nVETTED LIBRARY STORIES -- tailor each to this role and stage (do NOT reproduce verbatim):\n"
        ]
        for i, s in enumerate(library_seeds, 1):
            tags_str = ", ".join(s.get("tags", []))
            perf = s.get("_performance_signal", "")
            seed_lines.append(f"LIBRARY STORY {i} [{tags_str}]:")
            if perf:
                seed_lines.append(f"Performance: {perf}")
            seed_lines.append(f"Employer: {s.get('employer', '')} | {s.get('title', '')} | {s.get('dates', '')}")
            seed_lines.append(f"Situation: {s.get('situation', '')}")
            seed_lines.append(f"Task: {s.get('task', '')}")
            seed_lines.append(f"Action: {s.get('action', '')}")
            seed_lines.append(f"Result: {s.get('result', '')}")
            if s.get("if_probed"):
                seed_lines.append(f"If probed: {s.get('if_probed')}")
            seed_lines.append("")
        seed_lines.append(
            "For each seeded story: tailor to this role/stage. "
            "After 'STORY N --' heading, append '(library-seeded)'. "
            "If a Performance line is shown above, reproduce it verbatim on the next line.\n"
        )
        seed_block = "\n".join(seed_lines)

    return (
        f"Generate employer-attributed interview stories for a {profile['label']}.\n\n"
        f"CANDIDATE PROFILE (PII removed):\n{candidate_profile[:2500]}\n\n"
        f"RESUME SUBMITTED FOR THIS ROLE -- with employer context:\n{story_context[:3000]}\n\n"
        f"JOB DESCRIPTION:\n{jd[:2000]}\n\n"
        f"CRITICAL INSTRUCTIONS:\n"
        f"- Every story MUST be grounded in the bullets shown above\n"
        f"- Every story MUST include employer attribution "
        f"(\"During my time at [Employer] as [Title], [dates]...\")\n"
        f"- Do NOT invent metrics or outcomes\n\n"
        f"{_depth_instructions[profile['story_depth']]}\n\n"
        f"{_gap_instructions[profile['gap_behavior']]}\n\n"
        f"{seed_block}"
        f"Generate {profile['story_count']} stories. Use this format:\n\n"
        f"ROLE FIT ASSESSMENT:\n[{role_fit_instruction}]\n\n"
        f"KEY THEMES TO LEAD WITH:\n"
        f"Theme 1 -- [Name]: [1-2 sentences]\n"
        f"Theme 2 -- [Name]: [1-2 sentences]\n\n"
        f"STORY BANK:\n\n"
        f"STORY 1 -- [JD Requirement this addresses]:\n"
        f"Employer: [Company | Title | Dates]\n"
        f"Situation: [Context]\n"
        f"Task: [What needed to be done]\n"
        f"Action: [What YOU did -- first person]\n"
        f"Result: [Outcome -- qualitative acceptable]\n"
        f"If probed: [One additional sentence -- omit for headline depth]\n\n"
        f"[Continue for all stories in the {profile['story_count']} range]\n\n"
        f"LIKELY INTERVIEW QUESTIONS:\n"
        f"[5-8 questions likely to be asked, with one-line approach each]"
    )

def _build_gap_prompt(jd, gaps_section, candidate_profile, profile, library_seeds=None):
    """Build the Section 3 gap prep prompt, parameterized by stage profile."""
    peer_frame_block = ""
    if profile["gap_behavior"] == "full_peer":
        peer_frame_block = f"\n\n{profile['peer_frame_prompt']}"

    gap_depth_note = ""
    if profile["gap_behavior"] == "note":
        gap_depth_note = (
            "\nFor hiring manager stage: for each gap, include a brief note on how "
            "the gap might surface in a program context and how to address it proactively."
        )

    gap_seed_block = ""
    if library_seeds:
        seed_lines = [
            "\nVETTED GAP RESPONSES FROM LIBRARY -- tailor each to this role:\n"
        ]
        for s in library_seeds:
            tags_str = ", ".join(s.get("tags", []))
            perf = s.get("_performance_signal", "")
            seed_lines.append(f"GAP: {s.get('gap_label', '')} [{tags_str}]:")
            if perf:
                seed_lines.append(f"Performance: {perf}")
            seed_lines.append(f"Honest answer: {s.get('honest_answer', '')}")
            seed_lines.append(f"Bridge: {s.get('bridge', '')}")
            seed_lines.append(f"Redirect: {s.get('redirect', '')}")
            seed_lines.append("")
        seed_lines.append(
            "For each seeded gap: tailor to this role/stage. "
            "After 'GAP N --' heading, append '(library-seeded)'. "
            "If a Performance line is shown, reproduce it verbatim on the next line.\n"
        )
        gap_seed_block = "\n".join(seed_lines)

    return (
        f"You are doing a two-step gap analysis grounded strictly in the JD text and "
        f"candidate profile. Follow these steps exactly.\n\n"
        f"STEP 1 -- EXTRACT ALL JD REQUIREMENTS:\n"
        f"Read the FULL job description below -- including required qualifications, preferred "
        f"qualifications, responsibilities, and any other stated criteria. Extract two lists:\n"
        f"  REQUIRED: skills, experience, tools, or credentials explicitly marked as required\n"
        f"  PREFERRED: skills or experience explicitly marked as preferred, desired, or a plus\n\n"
        f"Do not infer requirements from job type, title, seniority, or industry norms.\n"
        f"Only use what the JD text directly states.\n\n"
        f"FULL JOB DESCRIPTION:\n{jd}\n\n"
        f"STEP 2 -- CROSS-REFERENCE AGAINST CANDIDATE PROFILE:\n"
        f"Compare your extracted lists against the candidate profile below. A gap is valid if:\n"
        f"  - HARD GAP: JD lists it as REQUIRED and it is absent from the candidate's experience\n"
        f"  - PREFERRED GAP: JD lists it as PREFERRED and absent -- flag as lower severity\n\n"
        f"GUARDRAIL – Confirmed Tools and Skills are NOT gaps:\n"
        f"Do not flag as a gap any skill, tool, or methodology that appears in the candidate's "
        f"Confirmed Tools or Confirmed Skills sections – even if the candidate's depth is limited. "
        f"Presence in Confirmed Tools or Skills means that gap does not exist.\n\n"
        f"GUARDRAIL – Redirect honesty:\n"
        f"In the Redirect field: pivot ONLY to experience explicitly stated in the candidate profile. "
        f"Never suggest the candidate claim experience they do not hold. "
        f"If the candidate has adjacent domain exposure (e.g., a short-term engagement in a related domain), "
        f"frame the redirect around that specific domain context – not around methodology depth they may not have.\n\n"
        f"Expect to find 3-5 gaps. If you find zero, re-examine preferred qualifications.\n\n"
        f"CANDIDATE CONFIRMED GAPS:\n{gaps_section[:1500]}\n\n"
        f"CANDIDATE FULL PROFILE:\n{candidate_profile[:2000]}\n"
        f"{gap_depth_note}\n\n"
        f"{gap_seed_block}"
        f"For each gap provide a direct, confident talking point -- not apologetic.\n\n"
        f"Format exactly as:\n\n"
        f"GAP 1 -- [Topic] [REQUIRED or PREFERRED]:\n"
        f"Gap: [What the JD states and why it is a gap]\n"
        f"Honest answer: [What to say -- confident, not apologetic]\n"
        f"Bridge: [Connection to actual experience]\n"
        f"Redirect: [Strength to pivot toward]\n\n"
        f"GAP 2 -- [Topic] [REQUIRED or PREFERRED]:\n"
        f"[same format]\n\n"
        f"GAP 3 -- [Topic] [REQUIRED or PREFERRED]:\n"
        f"[same format]\n\n"
        f"HARD QUESTIONS TO PREPARE FOR:\n"
        f"[5 questions that will probe these gaps, with one-sentence approach each]"
        f"{peer_frame_block}"
    )

# ==============================================
# SALARY EXTRACTION
# ==============================================

def extract_salary(jd_text):
    """Extract salary range from JD and estimate offer guidance."""
    patterns = [
        r'\$\s*([\d,]+(?:\.\d+)?)[kK]?\s*[-\u2013\u2014to]+\s*\$?\s*([\d,]+(?:\.\d+)?)[kK]?',
        r'([\d,]+(?:\.\d+)?)\s*[-\u2013\u2014to]+\s*([\d,]+(?:\.\d+)?)\s*/\s*(?:year|yr|annually)',
    ]
    low = high = None
    for pattern in patterns:
        matches = re.findall(pattern, jd_text, re.IGNORECASE)
        if matches:
            for match in matches:
                try:
                    v1 = float(match[0].replace(',', ''))
                    v2 = float(match[1].replace(',', ''))
                    if v1 < 1000:
                        v1 *= 1000
                    if v2 < 1000:
                        v2 *= 1000
                    if 80000 <= v1 <= 400000 and 80000 <= v2 <= 400000:
                        low, high = min(v1, v2), max(v1, v2)
                        break
                except Exception:
                    continue
            if low and high:
                break

    if not low or not high:
        return {
            'found': False,
            'text': 'Salary range not found in JD',
            'guidance': 'Research Glassdoor, Levels.fyi, and LinkedIn Salary for this role.'
        }

    midpoint = (low + high) / 2
    offer_low = low + (high - low) * 0.45
    offer_high = low + (high - low) * 0.65
    anchor = low + (high - low) * 0.72
    floor = low + (high - low) * 0.35
    anchor_rounded = round(anchor / 5000) * 5000

    return {
        'found': True,
        'text': f"${low:,.0f} \u2013 ${high:,.0f} (midpoint ${midpoint:,.0f})",
        'guidance': (
            f"Posted range: ${low:,.0f} \u2013 ${high:,.0f}\n"
            f"  Realistic offer zone: ${offer_low:,.0f} \u2013 ${offer_high:,.0f} "
            f"(companies rarely open at the top of range)\n"
            f"  Suggested anchor: ${anchor_rounded:,.0f} \u2013 already rounded for natural delivery\n"
            f"  If asked: '${anchor_rounded:,.0f} based on my 20+ years of defense SE "
            f"experience and current TS/SCI clearance, though I am open to "
            f"discussing total compensation.'\n"
            f"  Floor: Do not accept below ${floor:,.0f} for this role level."
        )
    }

# ==============================================
# JD TAG EXTRACTION
# ==============================================

def _extract_jd_tags(jd_text: str) -> list:
    tags = _ilp.load_tags()
    jd_lower = jd_text.lower()
    return [t for t in tags if t.lower() in jd_lower or t.lower().replace("-", " ") in jd_lower]


def _load_all_debriefs_safe():
    try:
        from scripts.phase5_debrief_utils import load_all_debriefs
        return load_all_debriefs()
    except Exception:
        return []


def _get_story_signal_safe(library_id, all_debriefs):
    try:
        from scripts.phase5_debrief_utils import get_story_performance_signal
        return get_story_performance_signal(library_id, all_debriefs)
    except Exception:
        return None


def _get_gap_signal_safe(gap_label, all_debriefs):
    try:
        from scripts.phase5_debrief_utils import get_gap_performance_signal
        return get_gap_performance_signal(gap_label, all_debriefs)
    except Exception:
        return None


def _load_role_debriefs_safe(role):
    try:
        from scripts.phase5_debrief_utils import load_debriefs
        return load_debriefs(role)
    except Exception:
        return []


def _load_salary_actuals_safe(debriefs):
    try:
        from scripts.phase5_debrief_utils import load_salary_actuals
        return load_salary_actuals(debriefs)
    except Exception:
        return None


def _build_continuity_safe(debriefs):
    try:
        from scripts.phase5_debrief_utils import build_continuity_section
        return build_continuity_section(debriefs)
    except Exception:
        return ""

# ==============================================
# LOAD RESUME BULLETS FROM STAGE FILE
# ==============================================

def load_resume_bullets(stage4_path, stage2_path):
    """
    Extract bullets and summary from stage4_final.txt (or stage2_approved.txt).
    Stage files are the source of truth -- the .docx is presentation only.
    Returns (resume_data dict, source filename).
    """
    path = source = None
    if os.path.exists(stage4_path):
        path, source = stage4_path, "stage4_final.txt"
    elif os.path.exists(stage2_path):
        path, source = stage2_path, "stage2_approved.txt"
    else:
        return None, None

    with open(path, encoding='utf-8') as f:
        content = f.read()

    resume_data = {'source': source, 'summary': '', 'employers': {}}
    current_section = None
    current_employer = None

    for line in content.split('\n'):
        stripped = line.strip()

        # Skip structural lines
        if (stripped.startswith('=') or stripped.startswith('STAGE') or
                stripped.startswith('Role:') or stripped.startswith('Generated:') or
                stripped.startswith('INSTRUCTIONS') or stripped.startswith('Save as') or
                stripped.startswith('END OF')):
            continue

        if stripped == '## PROFESSIONAL SUMMARY':
            current_section = 'summary'
            continue
        elif stripped == '## CORE COMPETENCIES':
            current_section = 'competencies'
            continue
        elif stripped.startswith('## ') and stripped not in [
                '## PROFESSIONAL SUMMARY', '## CORE COMPETENCIES']:
            current_employer = stripped[3:].strip()
            current_section = 'employer'
            if current_employer not in resume_data['employers']:
                resume_data['employers'][current_employer] = []
            continue

        if current_section == 'summary':
            if stripped and not stripped.startswith('['):
                resume_data['summary'] += (' ' if resume_data['summary'] else '') + stripped

        elif current_section == 'employer' and current_employer:
            if stripped.startswith('- ') and not stripped.startswith('['):
                bullet = stripped[2:].strip()
                bullet = re.sub(r'\s*\[Source:[^\]]*\]', '', bullet).strip()
                bullet = re.sub(r'\s*\[Theme:[^\]]*\]', '', bullet).strip()
                bullet = re.sub(r'\s*\[VERIFY[^\]]*\]', '', bullet).strip()
                if bullet:
                    resume_data['employers'][current_employer].append(bullet)

    return resume_data, source

# ==============================================
# BUILD EMPLOYER-ATTRIBUTED STORY CONTEXT
# ==============================================

def build_story_context(library, resume_data, jd_lower):
    """
    Match resume bullets to library entries to get employer attribution,
    dates, and additional context for STAR story building.
    """
    if not resume_data:
        return "No resume stage file found -- using candidate profile only."

    context_lines = []

    for emp_name, resume_bullets in resume_data['employers'].items():
        if not resume_bullets:
            continue

        # Find matching employer in library
        lib_employer = next(
            (e for e in library['employers']
             if e['name'].upper() == emp_name.upper() or
             emp_name.upper() in e['name'].upper()),
            None
        )

        context_lines.append(f"\n{'='*40}")
        context_lines.append(f"EMPLOYER: {emp_name}")

        if lib_employer:
            context_lines.append(f"Title: {lib_employer.get('title', '')}")
            context_lines.append(f"Dates: {lib_employer.get('dates', '')}")

        context_lines.append("Bullets on submitted resume (use these as story basis):")
        for b in resume_bullets:
            context_lines.append(f"  - {b}")

        # Add relevant library bullets for story depth
        if lib_employer:
            additional = []
            for lb in lib_employer['bullets']:
                if lb.get('flagged'):
                    continue
                kws = lb.get('keywords', [])
                if any(k.lower() in jd_lower for k in kws):
                    if not any(lb['text'][:60] in rb for rb in resume_bullets):
                        additional.append(
                            f"  - [{lb['theme']}] {lb['text'][:150]}"
                        )
            if additional:
                context_lines.append(
                    "Additional library context for story depth (not on resume):"
                )
                context_lines.extend(additional[:4])

    return "\n".join(context_lines)

# ==============================================
# DOCX GENERATION
# ==============================================

def generate_prep_docx(output_path, role, resume_source, stage_profile,
                        section1, section_intro, section2, section3, section4,
                        continuity_section=""):
    """
    Generate a clean formatted .docx interview prep document.
    Uses simple heading/normal/bullet styles -- no resume color scheme.
    """
    doc = Document()

    # Page margins
    from docx.shared import Inches
    sec = doc.sections[0]
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def add_normal(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)

    def add_bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    def parse_and_add_section(text):
        """
        Parse section text and add to doc with appropriate styles.
        Lines starting with all-caps followed by colon = Heading 3
        Lines starting with bullet markers = List Bullet
        Everything else = Normal
        """
        import re
        for line in text.split('\n'):
            stripped = line.strip()
            if not stripped:
                continue
            # Strip markdown bold markers left by the model
            stripped = re.sub(r'\*\*([^*]*)\*\*', r'\1', stripped)
            stripped = stripped.replace('**', '')
            # Strip markdown heading markers so they render as proper headings
            stripped = re.sub(r'^#{1,6}\s+', '', stripped)
            if not stripped:
                continue
            # All-caps heading pattern (e.g. "COMPANY OVERVIEW:", "STORY 1 -")
            if (re.match(r'^[A-Z][A-Z\s\-\u2013&/]+[:\-\u2013]', stripped) and
                    len(stripped) < 80 and not stripped.startswith('-')):
                add_heading(stripped, level=2)
            # Bullet lines
            elif stripped.startswith('- ') or stripped.startswith('\u2022 '):
                add_bullet(stripped.lstrip('-\u2022').strip())
            # Numbered items
            elif re.match(r'^\d+\.', stripped):
                add_bullet(stripped)
            # Story labels (Situation:, Task:, Action:, Result:)
            elif re.match(r'^(Situation|Task|Action|Result|Employer|'
                          r'Gap|Honest answer|Bridge|Redirect|Peer Frame|'
                          r'If probed|Theme \d|Follow-up):', stripped):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                # Label in bold
                label_end = stripped.index(':') + 1
                run1 = p.add_run(stripped[:label_end])
                run1.bold = True
                run2 = p.add_run(stripped[label_end:])
                run2.bold = False
            else:
                add_normal(stripped)

    # Title
    title_p = doc.add_heading('Interview Prep Package', 0)
    title_p.paragraph_format.space_after = Pt(4)

    # Metadata
    add_normal(f"Role: {role}")
    add_normal(f"Stage: {stage_profile['label']}")
    add_normal(f"Stage note: {stage_profile['description']}")
    add_normal(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    if resume_source:
        add_normal(f"Resume source: {resume_source}")
    doc.add_paragraph()

    # Section 1
    add_heading("Company & Role Brief", level=1)
    add_normal(f"({stage_profile['label']} -- verify currency before interview)")
    parse_and_add_section(section1)

    # Section 1.5
    add_heading("Introduce Yourself", level=1)
    add_normal(f"Tailored for {stage_profile['label']} register.")
    parse_and_add_section(section_intro)

    # Section 2
    add_heading("Story Bank", level=1)
    add_normal("Workshop stories before interview -- correct any overreach.")
    parse_and_add_section(section2)

    # Section 3
    add_heading("Gap Preparation", level=1)
    parse_and_add_section(section3)

    # Section 4
    add_heading("Questions to Ask", level=1)
    parse_and_add_section(section4)

    # Continuity section
    if continuity_section:
        add_heading("Continuity Summary", level=1)
        add_normal("(Reference record from prior interviews -- not prep guidance)")
        parse_and_add_section(continuity_section)

    doc.save(output_path)

# ==============================================
# WEB RESPONSE NORMALIZER
# ==============================================

def _normalize_web_response(text):
    """
    Clean up Section 1 text returned by the web_search tool.

    The tool leaves newlines where citation markers were stripped, producing
    sentences split across lines where the next fragment starts with punctuation
    (', with' / '. ') or a lowercase continuation word.  Also strips markdown
    heading markers (###) that the model sometimes emits despite plain-text
    instructions.
    """
    lines = text.split('\n')
    result = []
    for line in lines:
        clean = re.sub(r'^#{1,6}\s+', '', line)  # strip markdown heading markers
        stripped_clean = clean.strip()
        if not result:
            result.append(clean)
            continue
        prev = result[-1]
        # Join if this fragment starts with punctuation (citation-break artifact)
        if stripped_clean and stripped_clean[0] in ',.;:)':
            result[-1] = prev.rstrip() + stripped_clean
        # Join if previous line has no sentence-terminal punctuation and this
        # fragment continues mid-sentence (starts with a lowercase letter)
        elif (stripped_clean
              and prev.strip()
              and not prev.rstrip().endswith(('.', '!', '?', '”', '"', ':', ';'))
              and stripped_clean[0].islower()):
            result[-1] = prev.rstrip() + ' ' + stripped_clean
        else:
            result.append(clean)
    return '\n'.join(result)


# ==============================================
# CORE GENERATION FUNCTION
# ==============================================

def generate_prep(client, role_data, interview_stage, output_txt_path, output_docx_path,
                  dry_run=False):
    """
    Generate interview prep package from role data.
    role_data keys: jd_text, stage_text, library, candidate_profile, role_name.
    interview_stage: one of VALID_STAGES ('recruiter', 'hiring_manager', 'team_panel').
    dry_run: if True, print stage profile and return without API calls or file writes.
    Writes both .txt and .docx output files.
    All PII stripped from API payloads.
    """
    profile = STAGE_PROFILES[interview_stage]

    if dry_run:
        print("\nDRY RUN -- Stage profile that will be applied:")
        print(f"  Stage:       {profile['label']}")
        print(f"  Description: {profile['description']}")
        print(f"  Story count: {profile['story_count']}")
        print(f"  Story depth: {profile['story_depth']}")
        print(f"  Gap behavior:{profile['gap_behavior']}")
        print(f"  Salary in S1:{profile['salary_in_section1']}")
        print("\nNo API calls made. No files written.")
        return

    jd = role_data["jd_text"]
    raw_stage = role_data.get("stage_text", "")
    library = role_data["library"]
    raw_profile = role_data.get("candidate_profile", "")
    role_name = role_data.get("role_name", "unknown")

    # Strip PII from all text sent to API
    candidate_profile = strip_pii(raw_profile)
    safe_stage = strip_pii(raw_stage)
    jd_lower = jd.lower()

    # Extract JD tags for library filtering
    jd_tags = _extract_jd_tags(jd)

    # Load role debriefs for salary actuals and continuity (reused in Tasks 4 and 5)
    role_debriefs = _load_role_debriefs_safe(role_name)
    salary_actuals = _load_salary_actuals_safe(role_debriefs)

    # Build resume data from stage text (parse inline)
    resume_data = _parse_stage_text(safe_stage, source_label="stage_text")
    resume_source = resume_data.get('source') if resume_data else None

    story_context = build_story_context(library, resume_data, jd_lower)
    salary_data = extract_salary(jd)

    # Extract confirmed gaps section from profile
    gaps_section = ""
    if 'CONFIRMED GAPS' in raw_profile:
        start = raw_profile.find('CONFIRMED GAPS')
        end = raw_profile.find('## STYLE RULES', start)
        gaps_section = strip_pii(raw_profile[start:end if end > 0 else start + 2000])

    # --------------------------------------------------
    # SECTION 1 -- COMPANY & ROLE BRIEF (WEB-INFORMED)
    # --------------------------------------------------
    print("\nSection 1: Company & Role Brief (searching web)...")

    company_prompt = _build_section1_prompt(jd, salary_data, profile,
                                             salary_actuals=salary_actuals)

    response1 = client.messages.create(
        model=MODEL,
        max_tokens=1500,
        system=SYSTEM_PROMPT,
        tools=[{"type": "web_search_20250305", "name": "web_search"}],
        messages=[{"role": "user", "content": company_prompt}]
    )

    # Extract text blocks from response -- web search returns mixed content types
    section1_parts = []
    for block in response1.content:
        if hasattr(block, 'text') and block.text:
            section1_parts.append(block.text)
    section1 = _normalize_web_response(
        "\n".join(section1_parts) if section1_parts else
        "Web search unavailable -- review company website before interview."
    )

    # --------------------------------------------------
    # SECTION 1.5 -- INTRODUCE YOURSELF
    # --------------------------------------------------
    print("Section 1.5: Introduce Yourself (tailoring for stage)...")

    raw_intro = extract_profile_section(raw_profile, "INTRO MONOLOGUE")
    if raw_intro:
        intro_prompt = _build_intro_prompt(strip_pii(raw_intro), profile)
        response_intro = client.messages.create(
            model=MODEL,
            max_tokens=500,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": intro_prompt}]
        )
        section_intro = response_intro.content[0].text
    else:
        section_intro = (
            "No INTRO MONOLOGUE section found in candidate_profile.md. "
            "Add one to enable stage-tailored introduction generation."
        )

    # --------------------------------------------------
    # SECTION 2 -- STORY BANK (LIBRARY-GROUNDED)
    # --------------------------------------------------
    print("Section 2: Story Bank (grounded in resume and library)...")

    # Library seed lookup for stories
    story_seeds = _ilp.get_stories(tags=jd_tags) if jd_tags else []

    # Load debrief history once for both story and gap signal injection
    all_debriefs_for_signal = _load_all_debriefs_safe() if story_seeds else []

    # Inject performance signal into story seeds
    if story_seeds:
        for s in story_seeds:
            signal = _get_story_signal_safe(s.get("id"), all_debriefs_for_signal)
            if signal:
                s["_performance_signal"] = signal

    story_prompt = _build_section2_prompt(
        jd, story_context, candidate_profile, profile,
        library_seeds=story_seeds or None
    )

    response2 = client.messages.create(
        model=MODEL,
        max_tokens=3000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": story_prompt}]
    )
    section2 = response2.content[0].text

    # --------------------------------------------------
    # SECTION 3 -- GAP PREPARATION
    # --------------------------------------------------
    print("Section 3: Gap Preparation...")

    short_tenure_raw = extract_profile_section(raw_profile, "SHORT TENURE EXPLANATION")
    short_tenure_block = ""
    if short_tenure_raw:
        short_tenure_block = (
            "SHORT TENURE EXPLANATION:\n"
            + strip_pii(short_tenure_raw)
            + "\n\n" + "=" * 40 + "\n\n"
        )

    if profile["gap_behavior"] == "omit":
        section3 = (
            short_tenure_block
            + "Gap prep omitted -- do not volunteer gaps in a recruiter screen."
        )
    else:
        # Library seed lookup for gaps
        gap_seeds = _ilp.get_gap_responses(tags=jd_tags) if jd_tags else []

        # Inject performance signal into gap seeds (reuse already-loaded debrief history)
        if gap_seeds:
            if not all_debriefs_for_signal:
                all_debriefs_for_signal = _load_all_debriefs_safe()
            for g in gap_seeds:
                signal = _get_gap_signal_safe(g.get("gap_label"), all_debriefs_for_signal)
                if signal:
                    g["_performance_signal"] = signal

        gap_prompt = _build_gap_prompt(
            jd, gaps_section, candidate_profile, profile,
            library_seeds=gap_seeds or None
        )
        response3 = client.messages.create(
            model=MODEL,
            max_tokens=1200,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": gap_prompt}]
        )
        section3 = short_tenure_block + response3.content[0].text

    # --------------------------------------------------
    # SECTION 4 -- QUESTIONS TO ASK
    # --------------------------------------------------
    print("Section 4: Questions to Ask...")

    # Library seed lookup for questions
    question_seeds = _ilp.get_questions(tags=jd_tags, stage=interview_stage) if jd_tags else []
    question_seed_block = ""
    if question_seeds:
        qlines = ["\nVETTED QUESTIONS FROM LIBRARY -- include or adapt these as appropriate:\n"]
        for q in question_seeds:
            qlines.append(f"- {q.get('text', '')}")
        qlines.append(
            "\nFor each included library question, append '(library-seeded)' after the question text.\n"
        )
        question_seed_block = "\n".join(qlines)

    context_block = (
        f"JOB DESCRIPTION:\n{jd[:2000]}\n\n"
        f"CANDIDATE BACKGROUND (PII removed):\n{strip_pii(candidate_profile[:800])}\n\n"
        f"{question_seed_block}"
    )
    questions_prompt = context_block + profile["questions_prompt"]

    response4 = client.messages.create(
        model=MODEL,
        max_tokens=1200,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": questions_prompt}]
    )
    section4 = response4.content[0].text

    # Build continuity section from role debriefs
    continuity_text = _build_continuity_safe(role_debriefs)

    # --------------------------------------------------
    # COMPILE AND SAVE OUTPUT
    # --------------------------------------------------
    print("\nCompiling output...")

    output_lines = []
    output_lines.append("=" * 60)
    output_lines.append("INTERVIEW PREP PACKAGE v3")
    output_lines.append("=" * 60)
    output_lines.append(f"Role: {role_name}")
    output_lines.append(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    output_lines.append(f"Resume source: {resume_source if resume_source else 'Not found'}")
    output_lines.append(f"Stage: {profile['label']}")
    output_lines.append(f"Stage note: {profile['description']}")
    output_lines.append("Note: PII stripped from all API calls.")
    output_lines.append("=" * 60)
    output_lines.append("")

    output_lines.append("SECTION 1 \u2013 COMPANY & ROLE BRIEF")
    output_lines.append("(Web-informed -- verify currency before interview)")
    output_lines.append("-" * 60)
    output_lines.append(section1)
    output_lines.append("")

    output_lines.append("=" * 60)
    output_lines.append("SECTION 1.5 \u2013 INTRODUCE YOURSELF")
    output_lines.append(f"({profile['label']} register)")
    output_lines.append("-" * 60)
    output_lines.append(section_intro)
    output_lines.append("")

    output_lines.append("=" * 60)
    output_lines.append("SECTION 2 \u2013 STORY BANK")
    output_lines.append("(Grounded in submitted resume with employer attribution)")
    output_lines.append("REMINDER: Workshop stories in chat before interview.")
    output_lines.append("Correct any overreach -- every detail must be accurate.")
    output_lines.append("-" * 60)
    output_lines.append(section2)
    output_lines.append("")

    output_lines.append("=" * 60)
    output_lines.append("SECTION 3 \u2013 GAP PREPARATION")
    output_lines.append("-" * 60)
    output_lines.append(section3)
    output_lines.append("")

    output_lines.append("=" * 60)
    output_lines.append("SECTION 4 \u2013 QUESTIONS TO ASK")
    output_lines.append("-" * 60)
    output_lines.append(section4)
    output_lines.append("")

    if continuity_text:
        output_lines.append(continuity_text)
        output_lines.append("")

    output_lines.append("=" * 60)
    output_lines.append("END OF INTERVIEW PREP PACKAGE")
    output_lines.append(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    output_lines.append("=" * 60)

    output_text = "\n".join(output_lines)

    with open(output_txt_path, 'w', encoding='utf-8') as f:
        f.write(output_text)
    print(f"  Interview prep written to {output_txt_path}")

    # Generate docx
    try:
        generate_prep_docx(
            output_docx_path, role_name, resume_source, profile,
            section1, section_intro, section2, section3, section4,
            continuity_section=continuity_text
        )
        print(f"  Interview prep .docx written to {output_docx_path}")
    except Exception as e:
        print(f"  WARNING: Docx generation failed: {str(e)}")
        print(f"  Text file is still available: {output_txt_path}")
        # Still create a minimal docx so the file exists
        doc = Document()
        for line in output_text.splitlines():
            doc.add_paragraph(line)
        doc.save(output_docx_path)


def _parse_stage_text(stage_text, source_label="stage_text"):
    """
    Parse a stage file text string into resume_data dict.
    Mirrors load_resume_bullets() but operates on a string instead of a file path.
    """
    if not stage_text or not stage_text.strip():
        return None

    resume_data = {'source': source_label, 'summary': '', 'employers': {}}
    current_section = None
    current_employer = None

    for line in stage_text.split('\n'):
        stripped = line.strip()

        # Skip structural lines
        if (stripped.startswith('=') or stripped.startswith('STAGE') or
                stripped.startswith('Role:') or stripped.startswith('Generated:') or
                stripped.startswith('INSTRUCTIONS') or stripped.startswith('Save as') or
                stripped.startswith('END OF')):
            continue

        if stripped == '## PROFESSIONAL SUMMARY':
            current_section = 'summary'
            continue
        elif stripped == '## CORE COMPETENCIES':
            current_section = 'competencies'
            continue
        elif stripped.startswith('## ') and stripped not in [
                '## PROFESSIONAL SUMMARY', '## CORE COMPETENCIES']:
            current_employer = stripped[3:].strip()
            current_section = 'employer'
            if current_employer not in resume_data['employers']:
                resume_data['employers'][current_employer] = []
            continue

        if current_section == 'summary':
            if stripped and not stripped.startswith('['):
                resume_data['summary'] += (' ' if resume_data['summary'] else '') + stripped

        elif current_section == 'employer' and current_employer:
            if stripped.startswith('- ') and not stripped.startswith('['):
                bullet = stripped[2:].strip()
                bullet = re.sub(r'\s*\[Source:[^\]]*\]', '', bullet).strip()
                bullet = re.sub(r'\s*\[Theme:[^\]]*\]', '', bullet).strip()
                bullet = re.sub(r'\s*\[VERIFY[^\]]*\]', '', bullet).strip()
                if bullet:
                    resume_data['employers'][current_employer].append(bullet)

    return resume_data

# ==============================================
# MAIN
# ==============================================

def main():
    parser = argparse.ArgumentParser(description='Phase 5 Interview Prep Generator')
    parser.add_argument('--role', type=str, required=True,
                        help='Role package folder name (e.g. Acme_SE_Systems)')
    parser.add_argument('--interview_stage', type=str, default=None,
                        choices=VALID_STAGES,
                        help=f'Interview stage: {", ".join(VALID_STAGES)}')
    parser.add_argument('--dry_run', action='store_true',
                        help='Print stage profile and exit without generating output')
    args = parser.parse_args()

    role = args.role
    interview_stage = args.interview_stage

    # Interactive fallback if stage not provided
    if not interview_stage:
        print("\nSelect interview stage:")
        for i, s in enumerate(VALID_STAGES, 1):
            p = STAGE_PROFILES[s]
            print(f"  {i}. {s} – {p['label']}: {p['description']}")
        choice = input("Enter stage name or number (1-3): ").strip().lower()
        if choice in ("1", "recruiter"):
            interview_stage = "recruiter"
        elif choice in ("2", "hiring_manager"):
            interview_stage = "hiring_manager"
        elif choice in ("3", "team_panel"):
            interview_stage = "team_panel"
        else:
            print(f"Invalid selection '{choice}'. Valid stages: {', '.join(VALID_STAGES)}")
            sys.exit(1)

    package_dir = os.path.join(JOBS_PACKAGES_DIR, role)
    jd_path = os.path.join(package_dir, "job_description.txt")
    stage4_path = os.path.join(package_dir, "stage4_final.txt")
    stage2_path = os.path.join(package_dir, "stage2_approved.txt")
    output_txt_path = os.path.join(package_dir, f"interview_prep_{interview_stage}.txt")
    output_docx_path = os.path.join(package_dir, f"interview_prep_{interview_stage}.docx")

    print("=" * 60)
    print("PHASE 5 \u2013 INTERVIEW PREP GENERATOR v3")
    print("=" * 60)
    print(f"Role: {role}")
    print(f"Package: {package_dir}")

    errors = []
    if not os.path.exists(package_dir):
        errors.append(f"Job package folder not found: {package_dir}")
    if not os.path.exists(jd_path):
        errors.append(f"job_description.txt not found in {package_dir}")
    if not os.path.exists(CANDIDATE_PROFILE_PATH):
        errors.append(f"candidate_profile.md not found: {CANDIDATE_PROFILE_PATH}")
    if not os.path.exists(EXPERIENCE_LIBRARY):
        errors.append(f"experience_library.json not found: {EXPERIENCE_LIBRARY}")

    if errors:
        print("\nERRORS -- cannot proceed:")
        for e in errors:
            print(f"  {e}")
        sys.exit(1)

    print("\nLoading data...")

    with open(jd_path, encoding='utf-8') as f:
        jd_text = f.read()

    with open(CANDIDATE_PROFILE_PATH, encoding='utf-8') as f:
        candidate_profile = f.read()

    with open(EXPERIENCE_LIBRARY, encoding='utf-8') as f:
        library = json.load(f)
    print(f"  Experience library loaded: {library['metadata']['total_bullets']} bullets")

    # Load stage file text
    stage_text = ""
    if os.path.exists(stage4_path):
        with open(stage4_path, encoding='utf-8') as f:
            stage_text = f.read()
        print(f"  Resume loaded from stage4_final.txt")
    elif os.path.exists(stage2_path):
        with open(stage2_path, encoding='utf-8') as f:
            stage_text = f.read()
        print(f"  Resume loaded from stage2_approved.txt")
    else:
        print(f"  WARNING: No stage file found -- stories will use library only")

    # Overwrite protection
    if os.path.exists(output_txt_path):
        print(f"\nWARNING: interview_prep_{interview_stage}.txt already exists.")
        overwrite = input("  Overwrite? (y/n): ").strip().lower()
        if overwrite != 'y':
            print("  Cancelled. Existing file preserved.")
            sys.exit(0)

    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    role_data = {
        "jd_text": jd_text,
        "stage_text": stage_text,
        "library": library,
        "candidate_profile": candidate_profile,
        "role_name": role,
    }

    generate_prep(client, role_data, interview_stage, output_txt_path, output_docx_path,
                  dry_run=args.dry_run)

    print(f"\n{'=' * 60}")
    print("PHASE 5 COMPLETE")
    print(f"{'=' * 60}")
    print(f"Output saved: {output_txt_path}")
    print(f"\nNext steps:")
    print(f"  1. Open {output_docx_path} in Word for formatted reading")
    print(f"     Or open {output_txt_path} in VS Code (View > Word Wrap)")
    print(f"  2. Verify company brief accuracy (web results may be stale)")
    print(f"  3. Workshop STAR stories in chat -- correct any overreach")
    print(f"  4. Practice gap answers out loud -- confident, not apologetic")
    print(f"  5. Select 4-5 questions from Section 4")
    print(f"{'=' * 60}")

    # Post-generation notifications
    try:
        from scripts.phase5_debrief_utils import (
            load_debriefs, find_unmatched_debrief_content, has_debrief_for_stage
        )
        _notify_debriefs = load_debriefs(role)
        unmatched_stories, unmatched_gaps = find_unmatched_debrief_content(_notify_debriefs)
        if unmatched_stories or unmatched_gaps:
            print("\nDebrief content found that is not in your interview library.")
            print(f"Run: python -m scripts.phase5_workshop_capture --role {role} --stage {interview_stage}")
            print("to review and add workshopped content to the library.")

        if has_debrief_for_stage(_notify_debriefs, interview_stage):
            print("\nDebrief found for this stage. Generate thank you letters:")
            thankyou_cmd = f"python -m scripts.phase5_thankyou --role {role} --stage {interview_stage}"
            for d in _notify_debriefs:
                meta = d.get("metadata", {}) or {}
                if meta.get("stage") == interview_stage and meta.get("panel_label"):
                    thankyou_cmd += f" --panel_label {meta['panel_label']}"
                    break
            print(f"Run: {thankyou_cmd}")
    except Exception:
        pass  # notifications are best-effort; never block on failure


if __name__ == "__main__":
    main()
