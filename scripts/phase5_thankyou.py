# ==============================================
# phase5_thankyou.py
# Generates post-interview thank you letters
# for each interviewer from a filed debrief JSON.
#
# One .txt and one .docx per interviewer.
# Single Claude API call per letter.
# No subagents. No multi-turn.
#
# Usage:
#   python scripts/phase5_thankyou.py \
#     --role Viasat_SE_IS --stage hiring_manager
#   python scripts/phase5_thankyou.py \
#     --role Viasat_SE_IS --stage panel --panel_label se_team
# ==============================================

import os
import sys
import glob
import json
import argparse
from datetime import date
from anthropic import Anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches

# Add project root to path for utils import
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from scripts.utils.pii_filter import strip_pii

load_dotenv()

# ==============================================
# CONFIGURATION
# ==============================================

JOBS_PACKAGES_DIR = "data/job_packages"
DEBRIEFS_DIR = "data/debriefs"
CANDIDATE_PROFILE_PATH = "data/experience_library/candidate_profile.md"

MODEL = "claude-sonnet-4-20250514"

# ==============================================
# SYSTEM PROMPT
# ==============================================

SYSTEM_PROMPT = """You are an expert career coach who writes post-interview thank you letters
for senior professionals. Your letters are specific, confident, and brief -- never hollow.

You always:
- Open by referencing the specific content from the interviewer notes -- not a generic opener
- Write specific over generic, confident over effusive, brief over thorough
- Use en dashes, never em dashes
- Stay consistent with what the candidate said in the interview
- Reference stories and experience naturally -- do not recite them verbatim
- Keep letters to 3-4 paragraphs maximum

You never:
- Open with "Thank you for taking the time" or similar hollow openers
- Introduce new positions, claims, or experience not in the candidate profile or resume
- Reproduce verbatim STAR story text
- Use filler phrases or generic "I am excited about this opportunity" language"""

# ==============================================
# TONE CALIBRATION
# ==============================================

def _infer_tone(title):
    """
    Map interviewer title to a tone instruction string.
    Case-insensitive keyword match. Returns professional/warm default on no match.
    """
    if not title:
        return (
            "professional and warm -- process-aware, collegial, not overly technical"
        )

    t = title.lower()

    executive_keywords = [
        "director", "vp", "president", "bd", "manager", "pm", "program"
    ]
    technical_keywords = ["engineer", "se", "architect", "developer", "scientist"]
    recruiter_keywords = ["recruiter", "talent", "hr", "sourcer"]

    if any(k in t for k in executive_keywords):
        return (
            "mission outcomes and strategic framing -- lead with impact and "
            "organizational fit"
        )
    if any(k in t for k in technical_keywords):
        return (
            "peer-level technical -- write as one engineer to another; "
            "reference the specific work discussed"
        )
    if any(k in t for k in recruiter_keywords):
        return (
            "professional and warm -- process-aware, collegial, not overly technical"
        )

    return "professional and warm -- process-aware, collegial, not overly technical"


# ==============================================
# DEBRIEF FILE SELECTION
# ==============================================

def _find_debrief(role, stage, debriefs_dir, panel_label=None):
    """
    Glob for filed debrief JSON for the given role/stage/panel_label.
    Returns the path to the most recent match, or exits with a clear error.
    """
    role_dir = os.path.join(debriefs_dir, role)

    if panel_label:
        pattern = os.path.join(
            role_dir, f"debrief_{stage}_{panel_label}_*_filed-*.json"
        )
    else:
        # Only match unlabeled files: date starts immediately after stage
        pattern = os.path.join(
            role_dir, f"debrief_{stage}_[0-9]*_filed-*.json"
        )

    matches = sorted(glob.glob(pattern), reverse=True)

    if not matches:
        label_clause = f" with panel_label '{panel_label}'" if panel_label else ""
        print(
            f"\nERROR: No filed debrief found for role='{role}', stage='{stage}'"
            f"{label_clause}.\n"
            f"Expected pattern: {pattern}\n"
            f"Run phase5_debrief.py --file to produce a filed debrief first."
        )
        sys.exit(1)

    print(f"\nDebrief files found ({len(matches)}):")
    for m in matches:
        print(f"  {m}")
    selected = matches[0]
    print(f"Selected (most recent): {selected}")

    return selected


# ==============================================
# INPUT LOADING
# ==============================================

def _load_inputs(role, debrief_path):
    """
    Load JD, resume, candidate profile, and debrief JSON.
    Returns (inputs dict, warnings list).
    Exits on missing required files.
    """
    warnings = []
    package_dir = os.path.join(JOBS_PACKAGES_DIR, role)

    # Debrief JSON (already located)
    with open(debrief_path, encoding="utf-8") as f:
        debrief = json.load(f)

    # Job description -- required
    jd_path = os.path.join(package_dir, "job_description.txt")
    if not os.path.exists(jd_path):
        print(f"\nERROR: job_description.txt not found: {jd_path}")
        sys.exit(1)
    with open(jd_path, encoding="utf-8") as f:
        jd_text = f.read()

    # Resume -- optional, warn if missing
    resume_text = None
    stage4_path = os.path.join(package_dir, "stage4_final.txt")
    stage2_path = os.path.join(package_dir, "stage2_approved.txt")
    if os.path.exists(stage4_path):
        with open(stage4_path, encoding="utf-8") as f:
            resume_text = f.read()
    elif os.path.exists(stage2_path):
        with open(stage2_path, encoding="utf-8") as f:
            resume_text = f.read()
    else:
        warnings.append(
            f"WARNING: No resume file found (stage4_final.txt or stage2_approved.txt) "
            f"in {package_dir}. Letters will proceed without resume summary."
        )

    # Candidate profile -- required
    if not os.path.exists(CANDIDATE_PROFILE_PATH):
        print(f"\nERROR: candidate_profile.md not found: {CANDIDATE_PROFILE_PATH}")
        sys.exit(1)
    with open(CANDIDATE_PROFILE_PATH, encoding="utf-8") as f:
        candidate_profile = f.read()

    return {
        "debrief": debrief,
        "jd_text": jd_text,
        "resume_text": resume_text,
        "candidate_profile": candidate_profile,
    }, warnings


# ==============================================
# PROMPT ASSEMBLY
# ==============================================

def _build_letter_prompt(interviewer, index, debrief, jd_text, candidate_profile,
                          resume_text):
    """
    Assemble the user message for a single interviewer's thank you letter.
    All text sent to the API is passed through strip_pii().
    """
    name = interviewer.get("name") or f"Interviewer {index + 1}"
    title = interviewer.get("title") or ""
    notes = interviewer.get("notes") or ""

    tone = _infer_tone(title)

    stories_used = debrief.get("stories_used") or []
    landed_stories = [s for s in stories_used if s.get("landed") == "yes"]
    what_i_said = debrief.get("what_i_said") or ""
    gaps_surfaced = debrief.get("gaps_surfaced") or []

    # Build parts
    parts = []

    parts.append(f"INTERVIEWER:\nName: {strip_pii(name)}\nTitle: {strip_pii(title)}")

    if notes:
        parts.append(
            f"INTERVIEW NOTES (primary personalization anchor -- open by referencing this):\n"
            f"{strip_pii(notes)}"
        )
    else:
        parts.append(
            "INTERVIEW NOTES: None provided. Use the job description and candidate profile "
            "to anchor the opening."
        )

    if landed_stories:
        story_lines = []
        for s in landed_stories:
            framing = s.get("framing") or ""
            tags = s.get("tags") or []
            if framing or tags:
                story_lines.append(
                    f"  - {strip_pii(framing)}"
                    + (f" [tags: {', '.join(tags)}]" if tags else "")
                )
        if story_lines:
            parts.append(
                "STORIES THAT LANDED (reference naturally -- do not recite verbatim):\n"
                + "\n".join(story_lines)
            )

    if what_i_said:
        parts.append(
            f"CONTINUITY (stay consistent -- do not introduce new positions):\n"
            f"{strip_pii(what_i_said)}"
        )

    jd_excerpt = strip_pii(jd_text[:1500])
    parts.append(f"JOB DESCRIPTION EXCERPT:\n{jd_excerpt}")

    profile_excerpt = strip_pii(candidate_profile[:1000])
    parts.append(f"CANDIDATE PROFILE EXCERPT:\n{profile_excerpt}")

    if resume_text:
        resume_summary = strip_pii(resume_text[:800])
        parts.append(f"RESUME SUMMARY:\n{resume_summary}")

    parts.append(f"TONE: {tone}")

    # Letter structure instructions
    gap_labels = [g.get("gap_label", "") for g in gaps_surfaced if g.get("gap_label")]
    if gap_labels:
        gaps_block = (
            "LETTER STRUCTURE:\n"
            "Write 3-4 paragraphs:\n"
            "  P1: Open by referencing the specific content from the interview notes.\n"
            "  P2: Reinforce strongest fit signal using a landed story or key theme.\n"
            "  P3: Stay consistent with what_i_said; affirm interest and mutual fit.\n"
            "  P4 (optional): Consider briefly reframing one of these gaps "
            "if it came up naturally in the conversation and a reframe adds value: "
            + ", ".join(gap_labels)
            + ". Omit P4 if it would feel forced."
        )
    else:
        gap_labels_block = (
            "LETTER STRUCTURE:\n"
            "Write 3 paragraphs:\n"
            "  P1: Open by referencing the specific content from the interview notes.\n"
            "  P2: Reinforce strongest fit signal using a landed story or key theme.\n"
            "  P3: Stay consistent with what_i_said; affirm interest and mutual fit."
        )
        gaps_block = gap_labels_block

    parts.append(gaps_block)
    parts.append(
        "Write the letter now. Output only the letter body -- no subject line, "
        "no salutation header, no signature block. Begin directly with the opening sentence."
    )

    return "\n\n".join(parts)


# ==============================================
# OUTPUT FILENAME HELPERS
# ==============================================

def _interviewer_lastname(name, index):
    """Extract lowercase last name token, or fall back to 'interviewerN'."""
    if not name or not name.strip():
        return f"interviewer{index + 1}"
    return name.strip().split()[-1].lower()


def _output_paths(package_dir, stage, panel_label, lastname, run_date):
    """Return (txt_path, docx_path) for one interviewer."""
    if panel_label:
        stem = f"thankyou_{stage}_{panel_label}_{lastname}_{run_date}"
    else:
        stem = f"thankyou_{stage}_{lastname}_{run_date}"
    return (
        os.path.join(package_dir, f"{stem}.txt"),
        os.path.join(package_dir, f"{stem}.docx"),
    )


# ==============================================
# DOCX CONSTRUCTION
# ==============================================

def _write_docx(docx_path, letter_body, interviewer_name, interviewer_title,
                 role, stage, run_date):
    """
    Write a flat-prose thank you letter as a .docx document.
    Structure: title paragraph, metadata paragraph, blank, letter body paragraphs.
    Font matches phase5_interview_prep.py convention (Word default Normal style).
    """
    doc = Document()

    # Page setup: US Letter, 1-inch margins
    section = doc.sections[0]
    from docx.shared import Inches as _Inches
    section.page_width = int(8.5 * 914400)   # 8.5 inches in EMU
    section.page_height = int(11 * 914400)    # 11 inches in EMU
    section.left_margin = _Inches(1.0)
    section.right_margin = _Inches(1.0)
    section.top_margin = _Inches(1.0)
    section.bottom_margin = _Inches(1.0)

    # Title paragraph
    name_display = interviewer_name or "Interviewer"
    title_display = interviewer_title or ""
    title_text = f"Thank You \u2013 {name_display}"
    if title_display:
        title_text += f", {title_display}"

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Metadata paragraph
    meta_text = f"Role: {role}  |  Stage: {stage}  |  Date: {run_date}"
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(meta_text)
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    # Blank paragraph
    doc.add_paragraph()

    # Letter body -- each paragraph as a separate Paragraph element
    for para_text in letter_body.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    doc.save(docx_path)


# ==============================================
# LETTER GENERATION
# ==============================================

def generate_letters(client, role, stage, panel_label, inputs, run_date):
    """
    Iterate interviewers from debrief, call API, write .txt and .docx per letter.
    Returns (generated list, skipped list) of (name, txt_path, docx_path) tuples.
    """
    debrief = inputs["debrief"]
    jd_text = inputs["jd_text"]
    candidate_profile = inputs["candidate_profile"]
    resume_text = inputs["resume_text"]

    interviewers = (debrief.get("metadata") or {}).get("interviewers") or []
    package_dir = os.path.join(JOBS_PACKAGES_DIR, role)

    generated = []
    skipped = []

    for index, interviewer in enumerate(interviewers):
        name = interviewer.get("name") or f"Interviewer {index + 1}"
        notes = interviewer.get("notes")

        if not notes:
            print(
                f"\nWARNING: No notes for {name}. "
                "Letter will be generated with reduced personalization."
            )

        lastname = _interviewer_lastname(name, index)
        txt_path, docx_path = _output_paths(
            package_dir, stage, panel_label, lastname, run_date
        )

        # Overwrite protection
        if os.path.exists(txt_path):
            answer = input(
                f"{os.path.basename(txt_path)} already exists. Overwrite? (y/n): "
            ).strip().lower()
            if answer != "y":
                print(f"  Skipped: {name}")
                skipped.append((name, txt_path, docx_path))
                continue

        print(f"\nGenerating letter for {name} ({interviewer.get('title') or 'no title'})...")

        prompt = _build_letter_prompt(
            interviewer, index, debrief, jd_text, candidate_profile, resume_text
        )

        response = client.messages.create(
            model=MODEL,
            max_tokens=800,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        )
        letter_body = response.content[0].text

        # Write .txt
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(letter_body)
        print(f"  Written: {txt_path}")

        # Write .docx
        try:
            _write_docx(
                docx_path,
                letter_body,
                name,
                interviewer.get("title"),
                role,
                stage,
                run_date,
            )
            print(f"  Written: {docx_path}")
        except Exception as e:
            print(f"  WARNING: .docx generation failed for {name}: {e}")
            print(f"  Text file is still available: {txt_path}")

        generated.append((name, txt_path, docx_path))

    return generated, skipped


# ==============================================
# MAIN
# ==============================================

def main():
    parser = argparse.ArgumentParser(
        description="Phase 5 Thank You Letter Generator"
    )
    parser.add_argument(
        "--role", type=str, required=True,
        help="Role package folder name (e.g. Viasat_SE_IS)"
    )
    parser.add_argument(
        "--stage", type=str, required=True,
        help="Interview stage (matches debrief stage field, e.g. hiring_manager)"
    )
    parser.add_argument(
        "--panel_label", type=str, default=None,
        help="Panel label used in debrief filename (e.g. se_team)"
    )
    args = parser.parse_args()

    role = args.role
    stage = args.stage
    panel_label = args.panel_label

    print("=" * 60)
    print("PHASE 5 \u2013 THANK YOU LETTER GENERATOR")
    print("=" * 60)
    print(f"Role:        {role}")
    print(f"Stage:       {stage}")
    if panel_label:
        print(f"Panel label: {panel_label}")

    # Locate debrief
    debrief_path = _find_debrief(role, stage, DEBRIEFS_DIR, panel_label=panel_label)

    # Load inputs
    inputs, warnings = _load_inputs(role, debrief_path)

    if warnings:
        print()
        for w in warnings:
            print(w)

    run_date = str(date.today())

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        print("\nERROR: ANTHROPIC_API_KEY not set in environment.")
        sys.exit(1)

    client = Anthropic(api_key=api_key)

    generated, skipped = generate_letters(
        client, role, stage, panel_label, inputs, run_date
    )

    # Summary block
    print(f"\n{'=' * 60}")
    if generated:
        print(f"Generated {len(generated)} thank you letter(s):")
        for name, txt_path, docx_path in generated:
            print(f"  {txt_path}")
            print(f"  {docx_path}")
    else:
        print("No letters generated.")

    if skipped:
        print("\nSkipped:")
        for name, _, _ in skipped:
            print(f"  {name} \u2013 file exists, user declined overwrite")

    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
