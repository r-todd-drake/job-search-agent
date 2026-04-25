# ==============================================
# phase4_resume_generator.py
# AI-powered resume tailoring pipeline.
# Four-stage asynchronous workflow:
#
# Stage 1: Keyword + semantic bullet selection
#          Reads JD, queries experience library,
#          outputs stage1_draft.txt for review
#
# Stage 2: Manual review (no script needed)
#          Edit stage1_draft.txt in VS Code
#          Save as stage2_approved.txt
#
# Stage 3: Semantic review + wording suggestions
#          Reads stage2_approved.txt + JD,
#          outputs stage3_review.txt
#
# Stage 4: Document generation
#          Reads stage4_final.txt (or stage2_approved.txt),
#          generates .docx + runs check_resume.py
#
# Usage:
#   python -m scripts.phase4_resume_generator --stage 1 --role Acme_MBSE_Lead
#   python -m scripts.phase4_resume_generator --stage 3 --role Acme_MBSE_Lead
#   python -m scripts.phase4_resume_generator --stage 4 --role Acme_MBSE_Lead
# ==============================================

import os
import re
import json
import argparse
import hashlib
import subprocess
from datetime import datetime
from anthropic import Anthropic
from dotenv import load_dotenv
from scripts.utils.pii_filter import strip_pii
from scripts.utils import candidate_config
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import docx

load_dotenv()

from scripts.config import (
    JOBS_PACKAGES_DIR,
    EXPERIENCE_LIBRARY_JSON as EXPERIENCE_LIBRARY,
    CANDIDATE_PROFILE_PATH,
    RESUME_TEMPLATE,
    RESUMES_DIR,
    MODEL_SONNET,
)

# ==============================================
# CONFIGURATION
# ==============================================

CHECK_RESUME_SCRIPT = "scripts/check_resume.py"


def _get_employer_tiers():
    employers = candidate_config.load().get("employers", [])
    return {e["name"]: e["tier"] for e in employers}


def _get_chronological_order():
    employers = candidate_config.load().get("employers", [])
    return [e["name"] for e in employers]


# Max candidate bullets per employer sent to semantic selection
# More candidates = better selection but higher API cost
MAX_CANDIDATES_TIER1 = 12
MAX_CANDIDATES_TIER2 = 6
MAX_CANDIDATES_TIER3 = 4

# Load canonical candidate profile dynamically
# This replaces the hardcoded profile and prevents hallucinations
def load_candidate_profile():
    if os.path.exists(CANDIDATE_PROFILE_PATH):
        with open(CANDIDATE_PROFILE_PATH, encoding='utf-8') as f:
            return f.read()
    else:
        print(f"WARNING: candidate_profile.md not found at {CANDIDATE_PROFILE_PATH}")
        print("  Run phase3_build_candidate_profile.py to generate it.")
        print("  Using fallback minimal profile.")
        try:
            return candidate_config.build_known_facts()
        except FileNotFoundError:
            return (
                "\nCANDIDATE: [CANDIDATE]\n"
                "CLEARANCE: [see candidate_config.yaml]\n"
                "Run phase3_build_candidate_profile.py and ensure candidate_config.yaml exists.\n"
            )

# ==============================================
# VALIDATION HELPERS
# ==============================================

def validate_inputs(stage, package_dir, jd_path, stage1_path, stage2_path,
                    stage4_path):
    """Validate required files exist before running a stage."""
    errors = []

    # Always required
    if not os.path.exists(package_dir):
        errors.append(
            f"Job package folder not found: {package_dir}\n"
            f"  Create the folder and add job_description.txt before running Stage 1."
        )
        return errors  # No point checking further

    if not os.path.exists(jd_path):
        errors.append(
            f"job_description.txt not found in {package_dir}\n"
            f"  Paste the full job description into {jd_path} before running Stage 1."
        )

    if not os.path.exists(EXPERIENCE_LIBRARY):
        errors.append(
            f"experience_library.json not found: {EXPERIENCE_LIBRARY}\n"
            f"  Run phase3_compile_library.py to generate it."
        )

    if stage >= 3:
        if not os.path.exists(stage2_path):
            errors.append(
                f"stage2_approved.txt not found in {package_dir}\n"
                f"  Review {stage1_path} in VS Code, make your edits,\n"
                f"  and save as stage2_approved.txt before running Stage 3."
            )
        elif os.path.exists(stage1_path):
            # Warn if stage2 appears to be unedited copy of stage1
            with open(stage1_path) as f:
                s1_hash = hashlib.md5(f.read().encode()).hexdigest()
            with open(stage2_path) as f:
                s2_hash = hashlib.md5(f.read().encode()).hexdigest()
            if s1_hash == s2_hash:
                errors.append(
                    f"WARNING: stage2_approved.txt appears identical to stage1_draft.txt.\n"
                    f"  Did you forget to review and edit stage1_draft.txt?\n"
                    f"  If you intentionally approved without changes, rename to continue."
                )

    if stage == 4:
        # Use stage4_final.txt if it exists, otherwise fall back to stage2_approved.txt
        if not os.path.exists(stage4_path) and not os.path.exists(stage2_path):
            errors.append(
                f"Neither stage4_final.txt nor stage2_approved.txt found in {package_dir}\n"
                f"  Complete Stage 2 review before running Stage 4."
            )

    return errors

def check_overwrite(filepath, stage_name):
    """Warn if output file already exists."""
    if os.path.exists(filepath):
        print(f"\nWARNING: {stage_name} output already exists: {filepath}")
        response = input("  Overwrite? (y/n): ").strip().lower()
        if response != 'y':
            print("  Cancelled. Existing file preserved.")
            return False
    return True

# ==============================================
# LOAD DATA
# ==============================================

def load_jd(jd_path):
    with open(jd_path, encoding='utf-8') as f:
        return f.read()

def load_library():
    with open(EXPERIENCE_LIBRARY, encoding='utf-8') as f:
        return json.load(f)

def load_text(filepath):
    with open(filepath, encoding='utf-8') as f:
        return f.read()

# ==============================================
# STAGE 1 – BULLET SELECTION
# ==============================================

def keyword_score_bullet(bullet, jd_lower):
    """Score a bullet based on keyword overlap with JD."""
    if not bullet.get('keywords'):
        return 0
    score = 0
    for kw in bullet['keywords']:
        if kw.lower() in jd_lower:
            score += 1
    return score

def stage1_select_bullets(client, jd, library, candidate_profile):
    """
    Stage 1: Keyword pre-filter + semantic selection.
    Returns structured draft content.
    """
    print("\nStage 1: Selecting bullets from experience library...")
    jd_lower = jd.lower()

    # Step 1A - Keyword pre-filter per employer
    candidates_by_employer = {}
    for employer in library['employers']:
        name = employer['name']
        tier = _get_employer_tiers().get(name, 2)
        max_candidates = {1: MAX_CANDIDATES_TIER1,
                         2: MAX_CANDIDATES_TIER2,
                         3: MAX_CANDIDATES_TIER3}.get(tier, 6)

        # Score and sort bullets - priority bullets are always included
        priority = []
        scored = []
        for bullet in employer['bullets']:
            if bullet.get('flagged'):
                continue
            if bullet.get('priority'):
                priority.append(bullet)
            else:
                score = keyword_score_bullet(bullet, jd_lower)
                scored.append((score, bullet))

        scored.sort(key=lambda x: -x[0])

        # Fill remaining slots with top-scored non-priority bullets
        remaining_slots = max(0, max_candidates - len(priority))
        top_scored = [b for score, b in scored[:remaining_slots]
                      if score > 0 or len(scored) <= remaining_slots]
        top = priority + top_scored

        # Always include at least some bullets for Tier 1 employers
        if tier == 1 and len(top) == 0:
            top = priority + [b for _, b in scored[:3]]

        if top:
            candidates_by_employer[name] = {
                'tier': tier,
                'bullets': top
            }

    print(f"  Keyword pre-filter: {sum(len(v['bullets']) for v in candidates_by_employer.values())} candidate bullets across {len(candidates_by_employer)} employers")

    # Step 1B - Semantic selection via Claude API
    print("  Running semantic selection...")

    candidates_text = ""
    for emp_name, data in candidates_by_employer.items():
        candidates_text += f"\n## {emp_name} (Tier {data['tier']})\n"
        for i, bullet in enumerate(data['bullets'], 1):
            candidates_text += f"  [{i}] {bullet['text']}\n"
            candidates_text += f"      Theme: {bullet['theme']}\n"

    # Select summaries
    summaries_text = ""
    for i, s in enumerate(library['summaries'], 1):
        summaries_text += f"[{i}] Theme: {s['theme']}\n{s['text'][:200]}...\n\n"

    clean_profile = strip_pii(candidate_profile)
    clean_jd = strip_pii(jd)

    prompt = f"""You are selecting resume content for a job application.

CANDIDATE PROFILE:
{clean_profile}

JOB DESCRIPTION:
{clean_jd[:3000]}

CANDIDATE BULLETS BY EMPLOYER (pre-filtered by keyword relevance):
{candidates_text}

AVAILABLE SUMMARIES:
{summaries_text}

INSTRUCTIONS:
1. Select the best bullets for each employer based on fit with this specific JD.
   - Tier 1 employers (Saronic, KForce, Shield AI, G2 OPS): select 3-5 bullets each
   - Tier 2 employers (SAIC): select 1-3 bullets if relevant, 0 if not
   - Tier 3 employers (L3, Army): select 0-1 bullets only if directly relevant
   - If an employer has low relevance, select fewer or no bullets
   - DEDUPLICATION: Never select two bullets from the same employer that cover
     substantially the same topic or use similar language. If two bullets make
     the same point, select only the stronger, more specific one.

2. Select the best matching summary (by number), or flag if none are a strong fit.

3. Return your response in EXACTLY this format:

SUMMARY_SELECTION: [number or "NONE - suggest new"]
SUMMARY_REASON: [one sentence why]

EMPLOYER: [exact employer name]
BULLETS_SELECTED: [comma-separated bullet numbers]
SELECTION_REASON: [one sentence rationale]

EMPLOYER: [next employer]
...

Do not include any other text. Follow the format exactly."""

    response = client.messages.create(
        model=MODEL_SONNET,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    selection_text = response.content[0].text

    # Step 1C - Generate core competencies
    print("  Generating core competencies...")
    cfg = candidate_config.load()
    gaps_list = "\n  ".join(cfg.get("confirmed_gaps", []))
    not_held = ", ".join(cfg.get("confirmed_skills", {}).get("not_held", []))
    comp_prompt = f"""You are generating a Core Competencies section for a defense systems engineering resume.

CANDIDATE PROFILE:
{clean_profile}

JOB DESCRIPTION:
{clean_jd[:2000]}

Generate 5-7 core competency bullet lines for this specific role.
Each line should follow this format: "Category: specific skill 1, specific skill 2, specific skill 3"

Rules:
- Only include competencies the candidate genuinely has based on the profile above
- Mirror JD language where accurate
- Always include a Clearance & Certifications line last
- The Clearance & Certifications line must be EXACTLY:
  "Clearance & Certifications: Current TS/SCI | ICAgile Certified Professional"
- Never invent experience, credentials, or education not in the candidate profile
- NEVER include tools the candidate does not have: {not_held}
- If a JD requires a tool the candidate lacks, omit it entirely from competencies
- Confirmed gaps (never misrepresent): {gaps_list}
- En dashes only, never em dashes

Return ONLY the competency lines, one per line, no numbering, no extra text."""

    comp_response = client.messages.create(
        model=MODEL_SONNET,
        max_tokens=500,
        messages=[{"role": "user", "content": comp_prompt}]
    )

    competencies = comp_response.content[0].text.strip()

    # Step 1D - Parse selections and build draft
    draft = build_stage1_draft(jd, selection_text, candidates_by_employer,
                                library['summaries'], library, competencies)
    return draft

def build_stage1_draft(jd, selection_text, candidates_by_employer,
                        summaries, library, competencies=""):
    """Parse Claude's selection response and build the stage1 draft."""

    lines = selection_text.strip().split('\n')
    draft_lines = []

    # Header
    draft_lines.append("=" * 60)
    draft_lines.append("STAGE 1 DRAFT - FOR YOUR REVIEW")
    draft_lines.append("=" * 60)
    draft_lines.append(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    draft_lines.append("")
    draft_lines.append("INSTRUCTIONS:")
    draft_lines.append("  1. Review selected summary and bullets below")
    draft_lines.append("  2. Swap, add, or remove bullets as needed")
    draft_lines.append("  3. Adjust wording if desired")
    draft_lines.append("  4. Save this file as stage2_approved.txt when satisfied")
    draft_lines.append("  5. Run Stage 3 for semantic review before generating .docx")
    draft_lines.append("=" * 60)
    draft_lines.append("")

    # Parse summary selection
    summary_num = None
    summary_reason = ""
    current_employer = None
    employer_bullets = {}
    employer_reasons = {}

    for line in lines:
        line = line.strip()
        if line.startswith("SUMMARY_SELECTION:"):
            val = line.replace("SUMMARY_SELECTION:", "").strip()
            if "NONE" in val.upper():
                summary_num = None
            else:
                try:
                    summary_num = int(re.search(r'\d+', val).group())
                except:
                    summary_num = None
        elif line.startswith("SUMMARY_REASON:"):
            summary_reason = line.replace("SUMMARY_REASON:", "").strip()
        elif line.startswith("EMPLOYER:"):
            current_employer = line.replace("EMPLOYER:", "").strip()
            employer_bullets[current_employer] = []
        elif line.startswith("BULLETS_SELECTED:") and current_employer:
            nums_str = line.replace("BULLETS_SELECTED:", "").strip()
            try:
                nums = [int(n.strip()) for n in nums_str.split(',') if n.strip().isdigit()]
                employer_bullets[current_employer] = nums
            except:
                pass
        elif line.startswith("SELECTION_REASON:") and current_employer:
            employer_reasons[current_employer] = line.replace("SELECTION_REASON:", "").strip()

    # Write summary section
    draft_lines.append("## PROFESSIONAL SUMMARY")
    draft_lines.append("")
    if summary_num and summary_num <= len(summaries):
        selected_summary = summaries[summary_num - 1]
        draft_lines.append(f"[Source: Summary #{summary_num} - {selected_summary['theme']}]")
        draft_lines.append(f"[Reason: {summary_reason}]")
        draft_lines.append("")
        draft_lines.append(selected_summary['text'])
    else:
        draft_lines.append("[FLAG: No strong summary match found in library.]")
        draft_lines.append("[A suggested summary will be generated in Stage 3.]")
        draft_lines.append("[Placeholder - replace with approved summary text]")
    draft_lines.append("")

    # Write core competencies section
    draft_lines.append("## CORE COMPETENCIES")
    draft_lines.append("")
    if competencies:
        for line in competencies.split("\n"):
            line = line.strip()
            if line:
                draft_lines.append(f"- {line}")
    else:
        draft_lines.append("[No competencies generated - add manually]")
    draft_lines.append("")

    # Write employer sections in reverse chronological order
    chrono = _get_chronological_order()
    tier_order = sorted(
        candidates_by_employer.keys(),
        key=lambda x: chrono.index(x) if x in chrono else 99
    )

    for emp_name in tier_order:
        data = candidates_by_employer[emp_name]
        selected_nums = employer_bullets.get(emp_name, [])
        reason = employer_reasons.get(emp_name, "")

        # Get employer metadata
        emp_data = next((e for e in library['employers'] if e['name'] == emp_name), {})
        title = emp_data.get('title', '')
        dates = emp_data.get('dates', '')

        draft_lines.append(f"## {emp_name}")
        if title:
            draft_lines.append(f"Title: {title}")
        if dates:
            draft_lines.append(f"Dates: {dates}")
        if reason:
            draft_lines.append(f"[Selection rationale: {reason}]")
        draft_lines.append("")

        if not selected_nums:
            draft_lines.append("[No bullets selected - remove this section or add manually]")
        else:
            bullets = data['bullets']
            for num in selected_nums:
                if 1 <= num <= len(bullets):
                    bullet = bullets[num - 1]
                    draft_lines.append(f"- {bullet['text']}")
                    draft_lines.append(f"  [Source: {', '.join(bullet['sources'][:2])}]")
                    draft_lines.append(f"  [Theme: {bullet['theme']}]")
                    draft_lines.append("")

        draft_lines.append("")

    # Footer
    draft_lines.append("=" * 60)
    draft_lines.append("END OF STAGE 1 DRAFT")
    draft_lines.append("Save as stage2_approved.txt when ready for Stage 3.")
    draft_lines.append("=" * 60)

    return "\n".join(draft_lines)

# ==============================================
# EXTRACTED STAGE FUNCTIONS (testable)
# ==============================================

def run_stage1(client, jd_text, library, candidate_profile, output_path):
    """
    Run Stage 1: bullet selection and draft generation.
    Strips PII before API calls, writes draft to output_path.
    """
    draft = stage1_select_bullets(client, jd_text, library, candidate_profile)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(draft)


def run_stage3(client, stage2_text, jd_text, output_path):
    """
    Run Stage 3: semantic review of approved draft.
    Strips PII from both inputs before API call.
    Writes stage2_text + review notes to output_path.
    """
    clean_stage2 = strip_pii(stage2_text)
    clean_jd = strip_pii(jd_text)

    prompt = f"""You are reviewing a draft resume for a specific job application.
Your role is ADVISORY ONLY - suggest improvements, do not rewrite.

JOB DESCRIPTION:
{clean_jd[:3000]}

APPROVED RESUME DRAFT:
{clean_stage2}

Please provide your review in EXACTLY this format:

NARRATIVE COHERENCE:
[Does the summary claim align with the bullets selected? Are there coverage gaps?
Flag any mismatches between what the summary promises and what the bullets prove.]

COVERAGE GAPS:
[List any JD requirements not addressed by the current bullet selection.
For each gap, suggest whether a bullet exists that could fill it or if it should be acknowledged as a gap.]

WORDING SUGGESTIONS:
[For each suggested change, use this format:
ORIGINAL: [exact current text]
SUGGESTED: [proposed revision]
REASON: [why this improves fit, readability, or ATS alignment]
RULE CHECK: [confirm suggestion does not violate any candidate rules]

Only suggest changes grounded in confirmed candidate background.
Flag any suggestion that would require information not in the approved draft.]

ATS KEYWORDS:
[List JD keywords not currently appearing in the draft that could be
naturally incorporated without fabricating experience.]

SUMMARY ASSESSMENT:
[If the draft contains a library summary: confirm it is the best fit or suggest an alternative.
If the draft contains a placeholder: generate a suggested summary based on the selected bullets
and the style of existing summaries. Flag clearly as AI-GENERATED for review.]

OVERALL ASSESSMENT:
[1-2 sentences on overall fit and readiness for Stage 4.]"""

    response = client.messages.create(
        model=MODEL_SONNET,
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )

    review_text = response.content[0].text

    # Wrap in header
    output = []
    output.append("=" * 60)
    output.append("STAGE 3 SEMANTIC REVIEW - ADVISORY ONLY")
    output.append("=" * 60)
    output.append(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}")
    output.append("")
    output.append("INSTRUCTIONS:")
    output.append("  1. Review suggestions below")
    output.append("  2. Accept or reject each suggestion in stage2_approved.txt")
    output.append("  3. Save final approved content as stage4_final.txt")
    output.append("  4. Run Stage 4 to generate .docx")
    output.append("=" * 60)
    output.append("")
    output.append(review_text)
    output.append("")
    output.append("=" * 60)
    output.append("END OF STAGE 3 REVIEW")
    output.append("=" * 60)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(output))


def run_stage4(stage_text, output_path):
    """
    Run Stage 4: parse approved text and generate .docx resume.
    Pure docx generation - no API call.
    Falls back to basic Document() if template is not found.
    """
    sections = parse_final_content(stage_text)
    build_docx(sections, output_path)

# ==============================================
# STAGE 3 – SEMANTIC REVIEW (legacy wrapper)
# ==============================================

def stage3_semantic_review(client, jd, approved_content, role=""):
    """
    Stage 3: Semantic review of approved draft.
    Legacy wrapper - kept for main() compatibility.
    """
    print("\nStage 3: Running semantic review...")
    import tempfile, pathlib
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False,
                                     encoding='utf-8') as tmp:
        tmp_path = tmp.name
    run_stage3(client, approved_content, jd, tmp_path)
    with open(tmp_path, encoding='utf-8') as f:
        return f.read()

# ==============================================
# STAGE 4 – DOCUMENT GENERATION
# ==============================================

def stage4_generate_docx(final_content, role, resume_output_dir):
    """
    Stage 4: Parse final approved content and generate .docx resume.
    Uses python-docx with established resume formatting.
    """
    print("\nStage 4: Generating resume document...")

    os.makedirs(resume_output_dir, exist_ok=True)

    # Generate filename
    output_filename = f"{role}_Resume.docx"
    output_path = os.path.join(resume_output_dir, output_filename)

    run_stage4(final_content, output_path)

    return output_path

def parse_final_content(content):
    """Parse stage4_final.txt into structured sections for docx generation."""
    sections = {
        'summary': '',
        'competencies': [],
        'employers': []
    }

    current_section = None
    current_employer = None
    current_bullets = []
    current_title = ''
    current_dates = ''

    lines = content.split('\n')

    for line in lines:
        stripped = line.strip()

        # Skip header/footer lines
        if stripped.startswith('=') or stripped.startswith('STAGE') or \
           stripped.startswith('Role:') or stripped.startswith('Generated:') or \
           stripped.startswith('INSTRUCTIONS:') or stripped.startswith('  ') and \
           current_section is None:
            continue

        # Section detection
        if stripped == '## PROFESSIONAL SUMMARY':
            current_section = 'summary'
            continue

        if stripped == '## CORE COMPETENCIES':
            current_section = 'competencies'
            continue

        if stripped.startswith('## ') and stripped not in ['## PROFESSIONAL SUMMARY', '## CORE COMPETENCIES']:
            # Save previous employer
            if current_employer:
                sections['employers'].append({
                    'name': current_employer,
                    'title': current_title,
                    'dates': current_dates,
                    'bullets': [b for b in current_bullets if b]
                })

            current_employer = stripped[3:].strip()
            current_title = ''
            current_dates = ''
            current_bullets = []
            current_section = 'employer'
            continue

        # Content parsing
        if current_section == 'summary':
            # Skip source/reason tags
            if stripped.startswith('[') and stripped.endswith(']'):
                continue
            if stripped and not stripped.startswith('['):
                sections['summary'] += (' ' if sections['summary'] else '') + stripped

        elif current_section == 'competencies':
            if stripped.startswith('- ') and not stripped.startswith('['):
                comp_text = stripped[2:].strip()
                if comp_text and not comp_text.startswith('['):
                    sections['competencies'].append(comp_text)

        elif current_section == 'employer':
            if stripped.startswith('Title:'):
                current_title = re.sub(r'\[.*?\]', '', stripped.replace('Title:', '')).strip()
            elif stripped.startswith('Dates:'):
                current_dates = re.sub(r'\[.*?\]', '', stripped.replace('Dates:', '')).strip()
            elif stripped.startswith('- ') and not stripped.startswith('['):
                # Extract bullet text only, skip source/theme lines
                bullet_text = stripped[2:].strip()
                if bullet_text and not bullet_text.startswith('['):
                    # Strip any [VERIFY] or [FLAGGED] tags that leaked from library
                    bullet_text = re.sub(r'\s*\[VERIFY[^\]]*\]', '', bullet_text).strip()
                    bullet_text = re.sub(r'\s*\[FLAGGED[^\]]*\]', '', bullet_text).strip()
                    if bullet_text:
                        current_bullets.append(bullet_text)
            # Skip [Source:], [Theme:], [Selection rationale:] lines

    # Save final employer
    if current_employer:
        sections['employers'].append({
            'name': current_employer,
            'title': current_title,
            'dates': current_dates,
            'bullets': [b for b in current_bullets if b]
        })

    return sections

def build_docx(sections, output_path):
    """
    Build the .docx file using the resume template.
    Opens resume_template.docx which contains all custom styles
    (Title, Heading 1, Heading 3, List Bullet, Strong) with
    Aptos font and color scheme already defined.
    Falls back to basic formatting if template not found.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Load template or fall back to blank document
    if os.path.exists(RESUME_TEMPLATE):
        doc = Document(RESUME_TEMPLATE)
        print(f"  Using template: {RESUME_TEMPLATE}")
    else:
        print(f"  WARNING: Template not found at {RESUME_TEMPLATE}")
        print(f"  Save a blank styled document as {RESUME_TEMPLATE}")
        print(f"  Falling back to basic formatting...")
        doc = Document()

    def add_paragraph(text, style='Normal', space_before=0, space_after=4):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def add_name_header():
        # Name - Title style (use Normal as fallback if Title style absent)
        try:
            p = add_paragraph('', style='Title', space_after=0)
        except Exception:
            p = add_paragraph('', style='Normal', space_after=0)
        p.add_run(os.getenv('CANDIDATE_NAME', '[CANDIDATE]'))

        # Role title - Heading 1 style
        try:
            p2 = add_paragraph('', style='Heading 1', space_after=0)
        except Exception:
            p2 = add_paragraph('', style='Normal', space_after=0)
        cfg = candidate_config.load()
        p2.add_run(cfg["resume_defaults"]["role_title"])

        # Contact line - Normal 10pt
        p3 = add_paragraph('', style='Normal', space_after=6)
        location = os.getenv('CANDIDATE_LOCATION', 'San Diego, CA')
        phone = os.getenv('CANDIDATE_PHONE', '')
        email = os.getenv('CANDIDATE_EMAIL', '')
        linkedin = os.getenv('CANDIDATE_LINKEDIN', '')
        github = os.getenv('CANDIDATE_GITHUB', '')
        contact_parts = [p for p in [location, phone, email, linkedin, github] if p]
        run3 = p3.add_run(" | ".join(contact_parts))
        run3.font.size = Pt(10)

    def add_section_heading(text):
        try:
            p = add_paragraph('', style='Heading 1', space_before=8, space_after=2)
        except Exception:
            p = add_paragraph('', style='Normal', space_before=8, space_after=2)
        p.add_run(text)

    def add_employer_heading(name, title, dates):
        # Employer on one line as Heading 3
        try:
            p = add_paragraph('', style='Heading 3', space_before=6, space_after=0)
        except Exception:
            p = add_paragraph('', style='Normal', space_before=6, space_after=0)
        p.add_run(name)

        # Role title (bold) | Dates on one line
        p2 = add_paragraph('', style='Normal', space_before=0, space_after=2)
        if title:
            run1 = p2.add_run(title)
            run1.bold = True
            try:
                run1.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
            except Exception:
                pass
        if dates:
            run2 = p2.add_run(f"  |  {dates}")
            run2.bold = False

    def add_bullet(text):
        try:
            p = add_paragraph('', style='List Bullet', space_after=2)
        except Exception:
            p = add_paragraph('', style='Normal', space_after=2)
        p.add_run(text)

    def add_normal(text, size=11):
        p = add_paragraph('', style='Normal', space_after=4)
        run = p.add_run(text)
        run.font.size = Pt(size)

    # Build document content
    add_name_header()

    # Professional Summary
    if sections['summary']:
        add_section_heading("Professional Summary")
        add_normal(sections['summary'])

    # Core Competencies
    if sections.get('competencies'):
        add_section_heading("Core Competencies")
        for comp in sections['competencies']:
            add_bullet(comp)

    # Professional Experience
    if sections['employers']:
        add_section_heading("Professional Experience")
        for emp in sections['employers']:
            if not emp['bullets']:
                continue
            add_employer_heading(emp['name'], emp['title'], emp['dates'])
            for bullet_text in emp['bullets']:
                add_bullet(bullet_text)

    # Education & Certifications
    add_section_heading("Education & Certifications")
    cfg = candidate_config.load()
    add_normal(cfg["resume_defaults"]["education_line"])
    add_normal(cfg["resume_defaults"]["certifications_line"])

    doc.save(output_path)
    print(f"  Document saved: {output_path}")

# ==============================================
# MAIN
# ==============================================

def main():
    parser = argparse.ArgumentParser(description='Phase 4 Resume Generator')
    parser.add_argument('--stage', type=int, required=True, choices=[1, 3, 4],
                        help='Stage to run (1, 3, or 4)')
    parser.add_argument('--role', type=str, required=True,
                        help='Role package folder name (e.g. Acme_MBSE_Lead)')
    args = parser.parse_args()

    role = args.role
    stage = args.stage
    package_dir = os.path.join(JOBS_PACKAGES_DIR, role)
    resume_output_dir = os.path.join(RESUMES_DIR, role)

    # File paths
    jd_path = os.path.join(package_dir, "job_description.txt")
    stage1_path = os.path.join(package_dir, "stage1_draft.txt")
    stage2_path = os.path.join(package_dir, "stage2_approved.txt")
    stage3_path = os.path.join(package_dir, "stage3_review.txt")
    stage4_path = os.path.join(package_dir, "stage4_final.txt")

    candidate_profile = strip_pii(load_candidate_profile())
    print("  Candidate profile loaded and PII stripped.")

    print("=" * 60)
    print(f"PHASE 4 - RESUME GENERATOR - STAGE {stage}")
    print("=" * 60)
    print(f"Role: {role}")
    print(f"Package: {package_dir}")

    # Validate inputs
    errors = validate_inputs(stage, package_dir, jd_path, stage1_path,
                             stage2_path, stage4_path)
    if errors:
        print("\nERRORS - cannot proceed:")
        for error in errors:
            print(f"\n  {error}")
        exit(1)

    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    # -- STAGE 1 --
    if stage == 1:
        if not check_overwrite(stage1_path, "Stage 1 draft"):
            exit(0)

        jd = load_jd(jd_path)
        library = load_library()

        run_stage1(client, jd, library, candidate_profile, stage1_path)

        print(f"\nStage 1 complete.")
        print(f"  Draft saved: {stage1_path}")
        print(f"\nNext steps:")
        print(f"  1. Open {stage1_path} in VS Code")
        print(f"  2. Review selected bullets and summary")
        print(f"  3. Make any edits")
        print(f"  4. Save as {stage2_path}")
        print(f"  5. Run: python -m scripts.phase4_resume_generator --stage 3 --role {role}")

    # -- STAGE 3 --
    elif stage == 3:
        if not check_overwrite(stage3_path, "Stage 3 review"):
            exit(0)

        jd = load_jd(jd_path)
        approved = load_text(stage2_path)

        run_stage3(client, approved, jd, stage3_path)

        print(f"\nStage 3 complete.")
        print(f"  Review saved: {stage3_path}")
        print(f"\nNext steps:")
        print(f"  1. Open {stage3_path} in VS Code")
        print(f"  2. Review suggestions - accept or reject each one")
        print(f"  3. Apply accepted changes to {stage2_path}")
        print(f"  4. Save final version as {stage4_path}")
        print(f"  5. Run: python -m scripts.phase4_resume_generator --stage 4 --role {role}")

    # -- STAGE 4 --
    elif stage == 4:
        # Use stage4_final.txt if exists, otherwise stage2_approved.txt
        input_path = stage4_path if os.path.exists(stage4_path) else stage2_path
        print(f"  Using input: {input_path}")

        docx_output = os.path.join(resume_output_dir, f"{role}_Resume.docx")
        if not check_overwrite(docx_output, "Resume .docx"):
            exit(0)

        final_content = load_text(input_path)
        output_path = stage4_generate_docx(final_content, role, resume_output_dir)

        # Run check_resume.py automatically
        print(f"\n  Running quality check...")
        try:
            check_path = output_path.replace(os.sep, '/')
            result = subprocess.run(
                ["python", CHECK_RESUME_SCRIPT, check_path],
                capture_output=True, text=True, timeout=30
            )
            if result.stdout:
                print(result.stdout)
            if result.stderr:
                print("  Check stderr:", result.stderr[:200])
            if result.returncode != 0:
                print("  WARNING: Quality check found errors - review before submitting.")
            else:
                print("  Quality check passed.")
        except FileNotFoundError:
            print(f"  WARNING: check_resume.py not found at {CHECK_RESUME_SCRIPT}")
            print("  Run manually: python -m scripts.check_resume " + output_path)
        except Exception as e:
            print(f"  WARNING: Quality check failed to run: {str(e)}")
            print("  Run manually: python -m scripts.check_resume " + output_path)

        print(f"\nStage 4 complete.")
        print(f"  Resume: {output_path}")
        print(f"\nNext steps:")
        print(f"  1. Open {output_path} in Word")
        print(f"  2. Review formatting and adjust if needed")
        print(f"  3. Export to PDF for submission")
        print(f"  4. Update jobs.csv and tracker with application details")


if __name__ == '__main__':
    main()
