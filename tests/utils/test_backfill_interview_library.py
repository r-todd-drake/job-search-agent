# tests/utils/test_backfill_interview_library.py
import json
import os
import sys
import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
import scripts.utils.backfill_interview_library as backfill
import scripts.interview_library_parser as ilp


# ── Shared helpers ────────────────────────────────────────────────────────────

def _seed_library(tmp_path, data, monkeypatch):
    lib_path = tmp_path / "interview_library.json"
    lib_path.write_text(json.dumps(data), encoding="utf-8")
    monkeypatch.setattr(ilp, "LIBRARY_PATH", str(lib_path))
    monkeypatch.setattr(backfill, "LIBRARY_PATH", str(lib_path))
    return lib_path


def _paras(texts, style="Normal"):
    return [(t, style, False) for t in texts]


def _make_full_docx(tmp_path, filename="interview_prep_hiring_manager.docx", with_gaps=True):
    """Build a minimal .docx with all four workshopped sections."""
    from docx import Document
    doc = Document()
    doc.add_heading("Introduce Yourself", level=1)
    doc.add_paragraph("I am a systems engineer with 10 years of experience.")
    doc.add_heading("Story Bank", level=1)
    doc.add_paragraph("STORY 1 – Leadership Story:")
    doc.add_paragraph("Employer: ACME Corp | Systems Engineer | 2021-2023")
    doc.add_paragraph("Situation: Sat.")
    doc.add_paragraph("Task: Task.")
    doc.add_paragraph("Action: Act.")
    doc.add_paragraph("Result: Res.")
    if with_gaps:
        doc.add_heading("Gap Preparation", level=1)
        doc.add_paragraph("GAP 1 – IP Networking [REQUIRED]:")
        doc.add_paragraph("Honest answer: I have not worked in IP networking.")
        doc.add_paragraph("Bridge: Adjacent experience with systems interfaces.")
        doc.add_paragraph("Redirect: Strong systems engineering foundation.")
    doc.add_heading("Questions to Ask", level=1)
    doc.add_paragraph("1. What does success look like at 6 months?")
    path = tmp_path / filename
    doc.save(str(path))
    return str(path)


# ── discover_docx_files ───────────────────────────────────────────────────────

def test_discover_finds_docx_in_role_subdirs(tmp_path):
    role_dir = tmp_path / "Acme_SE"
    role_dir.mkdir()
    (role_dir / "interview_prep_hiring_manager.docx").write_bytes(b"")
    (role_dir / "interview_prep_recruiter_screen.docx").write_bytes(b"")
    results = backfill.discover_docx_files(str(tmp_path))
    assert len(results) == 2
    roles = {r for _, r, _ in results}
    stages = {s for _, _, s in results}
    assert roles == {"Acme_SE"}
    assert "hiring_manager" in stages
    assert "recruiter_screen" in stages


def test_discover_skips_unrecognized_stage_filenames(tmp_path):
    role_dir = tmp_path / "Acme_SE"
    role_dir.mkdir()
    (role_dir / "interview_prep_custom_stage.docx").write_bytes(b"")
    results = backfill.discover_docx_files(str(tmp_path))
    assert results == []


def test_discover_skips_non_docx_files(tmp_path):
    role_dir = tmp_path / "Acme_SE"
    role_dir.mkdir()
    (role_dir / "interview_prep_hiring_manager.txt").write_bytes(b"")
    results = backfill.discover_docx_files(str(tmp_path))
    assert results == []


def test_discover_raises_when_base_dir_absent(tmp_path):
    with pytest.raises(FileNotFoundError):
        backfill.discover_docx_files(str(tmp_path / "nonexistent"))


def test_discover_returns_sorted_results(tmp_path):
    for role in ("Zeta_Role", "Alpha_Role"):
        d = tmp_path / role
        d.mkdir()
        (d / "interview_prep_hiring_manager.docx").write_bytes(b"")
    results = backfill.discover_docx_files(str(tmp_path))
    paths = [p for p, _, _ in results]
    assert paths == sorted(paths)


# ── _split_sections_backfill ──────────────────────────────────────────────────

def test_split_backfill_captures_introduce_yourself():
    paras = (
        _paras(["Introduce Yourself"], style="Heading 1") +
        _paras(["Start with your systems engineering background."]) +
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 – MBSE:"])
    )
    sections = backfill._split_sections_backfill(paras)
    assert any("background" in t for t, _, _ in sections["introduce_yourself"])


def test_split_backfill_intro_does_not_bleed_into_story_bank():
    paras = (
        _paras(["Introduce Yourself"], style="Heading 1") +
        _paras(["Intro content here."]) +
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 – Topic:"])
    )
    sections = backfill._split_sections_backfill(paras)
    assert not any("STORY" in t for t, _, _ in sections["introduce_yourself"])
    assert any("STORY" in t for t, _, _ in sections["story_bank"])


def test_split_backfill_captures_all_four_sections():
    paras = (
        _paras(["Introduce Yourself"], style="Heading 1") +
        _paras(["Intro text."]) +
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 – Topic:"]) +
        _paras(["Gap Preparation"], style="Heading 1") +
        _paras(["GAP 1 – Topic [REQUIRED]:"]) +
        _paras(["Questions to Ask"], style="Heading 1") +
        _paras(["1. What does success look like at 6 months?"])
    )
    sections = backfill._split_sections_backfill(paras)
    assert len(sections["introduce_yourself"]) == 1
    assert len(sections["story_bank"]) == 1
    assert len(sections["gap_prep"]) == 1
    assert len(sections["questions"]) == 1


def test_split_backfill_company_role_brief_discarded():
    paras = (
        _paras(["Company Role Brief"], style="Heading 1") +
        _paras(["Company overview."]) +
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 – Topic:"])
    )
    sections = backfill._split_sections_backfill(paras)
    assert not any("overview" in t for t, _, _ in sections["story_bank"])
    assert not any("overview" in t for t, _, _ in sections["introduce_yourself"])


# ── Intro helpers ─────────────────────────────────────────────────────────────

def test_extract_intro_text_joins_non_italic_lines():
    paras = [
        ("Line one.", "Normal", False),
        ("Line two.", "Normal", False),
    ]
    assert backfill._extract_intro_text(paras) == "Line one.\nLine two."


def test_extract_intro_text_skips_italic():
    paras = [
        ("Line one.", "Normal", False),
        ("Coaching note.", "Normal", True),
        ("Line two.", "Normal", False),
    ]
    assert backfill._extract_intro_text(paras) == "Line one.\nLine two."


def test_build_intro_entry_schema():
    entry = backfill._build_intro_entry(
        "I am a systems engineer.", "Acme_SE", "hiring_manager", "2026-05-06"
    )
    assert entry == {
        "id": "intro-Acme_SE-hiring_manager",
        "role": "Acme_SE",
        "stage": "hiring_manager",
        "text": "I am a systems engineer.",
        "last_updated": "2026-05-06",
    }


def test_find_duplicate_intro_matches_role_and_stage():
    library = {
        "introductions": [
            {"id": "intro-Acme_SE-hiring_manager", "role": "Acme_SE", "stage": "hiring_manager"}
        ],
        "stories": [], "gap_responses": [], "questions": [],
    }
    result = backfill._find_duplicate_intro(library, "Acme_SE", "hiring_manager")
    assert result is not None
    assert result["id"] == "intro-Acme_SE-hiring_manager"


def test_find_duplicate_intro_no_match_different_stage():
    library = {
        "introductions": [
            {"id": "intro-Acme_SE-hiring_manager", "role": "Acme_SE", "stage": "hiring_manager"}
        ],
        "stories": [], "gap_responses": [], "questions": [],
    }
    assert backfill._find_duplicate_intro(library, "Acme_SE", "recruiter_screen") is None


def test_find_duplicate_intro_returns_none_when_key_absent():
    library = {"stories": [], "gap_responses": [], "questions": []}
    assert backfill._find_duplicate_intro(library, "Acme_SE", "hiring_manager") is None


# ── _make_unique_id ───────────────────────────────────────────────────────────

def test_make_unique_id_returns_base_when_no_collision():
    assert backfill._make_unique_id("mbse-story", set()) == "mbse-story"


def test_make_unique_id_appends_2_on_first_collision():
    assert backfill._make_unique_id("mbse-story", {"mbse-story"}) == "mbse-story-2"


def test_make_unique_id_increments_until_clear():
    existing = {"mbse-story", "mbse-story-2", "mbse-story-3"}
    assert backfill._make_unique_id("mbse-story", existing) == "mbse-story-4"


def test_make_unique_id_respects_60_char_limit():
    long_id = "a" * 58
    result = backfill._make_unique_id(long_id, {long_id})
    assert len(result) <= 60


# ── _process_file ─────────────────────────────────────────────────────────────

def test_process_file_writes_all_entry_types(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    library = {"stories": [], "gap_responses": [], "questions": []}
    docx_path = _make_full_docx(tmp_path)
    log_lines = []
    stats = backfill._process_file(
        docx_path, "Acme_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    assert stats["written"] == 4   # 1 intro + 1 story + 1 gap + 1 question
    assert stats["skipped"] == 0
    assert len(library.get("introductions", [])) == 1
    assert len(library["stories"]) == 1
    assert len(library["gap_responses"]) == 1
    assert len(library["questions"]) == 1


def test_process_file_skips_duplicate_story(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    library = {
        "stories": [
            {"id": "acme-corp-story", "employer": "ACME Corp", "tags": [], "roles_used": ["OtherRole"]}
        ],
        "gap_responses": [], "questions": [],
    }
    docx_path = _make_full_docx(tmp_path)
    log_lines = []
    stats = backfill._process_file(
        docx_path, "Acme_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    assert stats["skipped"] >= 1
    assert len(library["stories"]) == 1
    assert any("SKIP story" in line for line in log_lines)


def test_process_file_skips_gap_for_recruiter_screen(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    library = {"stories": [], "gap_responses": [], "questions": []}
    docx_path = _make_full_docx(tmp_path, filename="interview_prep_recruiter_screen.docx",
                                 with_gaps=False)
    log_lines = []
    backfill._process_file(
        docx_path, "Acme_SE", "recruiter_screen", library, "2026-05-06", log_lines
    )
    assert len(library["gap_responses"]) == 0
    assert any("recruiter_screen" in line for line in log_lines)


def test_process_file_returns_warning_on_unreadable_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    bad_path = str(tmp_path / "bad.docx")
    with open(bad_path, "w") as f:
        f.write("not a docx")
    library = {"stories": [], "gap_responses": [], "questions": []}
    log_lines = []
    stats = backfill._process_file(
        bad_path, "Acme_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    assert len(stats["warnings"]) == 1
    assert stats["written"] == 0


def test_process_file_collision_generates_unique_id(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    library = {"stories": [], "gap_responses": [], "questions": []}
    docx_path = _make_full_docx(tmp_path)
    log_lines = []
    backfill._process_file(
        docx_path, "Acme_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    initial_count = len(library.get("introductions", []))
    docx_path2 = _make_full_docx(tmp_path, filename="interview_prep_hiring_manager2.docx")
    backfill._process_file(
        docx_path2, "Beta_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    assert len(library.get("introductions", [])) == initial_count + 1


# ── introductions[] schema extension ─────────────────────────────────────────

def test_introductions_key_created_when_absent(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    library = {"stories": [], "gap_responses": [], "questions": []}
    docx_path = _make_full_docx(tmp_path)
    log_lines = []
    backfill._process_file(
        docx_path, "Acme_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    assert "introductions" in library
    assert isinstance(library["introductions"], list)
    assert len(library["introductions"]) == 1


def test_introductions_appends_when_key_exists(tmp_path, monkeypatch):
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    library = {
        "stories": [], "gap_responses": [], "questions": [],
        "introductions": [{"id": "intro-Other_SE-hiring_manager", "role": "Other_SE",
                           "stage": "hiring_manager"}],
    }
    docx_path = _make_full_docx(tmp_path)
    log_lines = []
    backfill._process_file(
        docx_path, "Acme_SE", "hiring_manager", library, "2026-05-06", log_lines
    )
    assert len(library["introductions"]) == 2


# ── run_backfill ──────────────────────────────────────────────────────────────

def test_run_backfill_writes_library_and_log(tmp_path, monkeypatch):
    role_dir = tmp_path / "workshopped" / "Acme_SE"
    role_dir.mkdir(parents=True)
    _make_full_docx(tmp_path / "workshopped" / "Acme_SE")

    lib_path = tmp_path / "interview_library.json"
    lib_path.write_text(json.dumps({"stories": [], "gap_responses": [], "questions": []}))
    out_dir = tmp_path / "outputs"
    out_dir.mkdir()

    monkeypatch.setattr(ilp, "LIBRARY_PATH", str(lib_path))
    monkeypatch.setattr(backfill, "LIBRARY_PATH", str(lib_path))
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    monkeypatch.setattr(backfill, "OUTPUTS_DIR", str(out_dir))

    result = backfill.run_backfill(str(tmp_path / "workshopped"))

    assert result["written"] > 0
    assert result["files"] == 1
    saved = json.loads(lib_path.read_text())
    assert len(saved["stories"]) > 0
    log_files = list(out_dir.glob("library_backfill_*.txt"))
    assert len(log_files) == 1


def test_run_backfill_aborts_on_high_dup_ratio(tmp_path, monkeypatch):
    role_dir = tmp_path / "workshopped" / "Acme_SE"
    role_dir.mkdir(parents=True)
    _make_full_docx(tmp_path / "workshopped" / "Acme_SE")

    existing = {
        "stories": [
            {"id": "acme-corp-story", "employer": "ACME Corp", "tags": [], "roles_used": []}
        ],
        "gap_responses": [
            {"id": "ip-networking", "gap_label": "IP Networking", "roles_used": []}
        ],
        "questions": [
            {"id": "what-does-success", "text": "What does success look like at 6 months?",
             "roles_used": []}
        ],
        "introductions": [
            {"id": "intro-Acme_SE-hiring_manager", "role": "Acme_SE", "stage": "hiring_manager"}
        ],
    }
    lib_path = tmp_path / "interview_library.json"
    lib_path.write_text(json.dumps(existing))
    out_dir = tmp_path / "outputs"
    out_dir.mkdir()

    monkeypatch.setattr(ilp, "LIBRARY_PATH", str(lib_path))
    monkeypatch.setattr(backfill, "LIBRARY_PATH", str(lib_path))
    monkeypatch.setattr(ilp, "TAGS_PATH", str(tmp_path / "missing_tags.json"))
    monkeypatch.setattr(backfill, "OUTPUTS_DIR", str(out_dir))

    result = backfill.run_backfill(str(tmp_path / "workshopped"))

    assert result.get("aborted") is True
    saved = json.loads(lib_path.read_text())
    assert len(saved["stories"]) == 1
