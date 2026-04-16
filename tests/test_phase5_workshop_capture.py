import pytest
import json
import os
import sys
from unittest.mock import MagicMock, patch
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import scripts.phase5_workshop_capture as wc


# ── Helpers ───────────────────────────────────────────────────────────────────

def _seed_library(tmp_path, data, monkeypatch):
    """Write library data to a temp file and monkeypatch LIBRARY_PATH."""
    import scripts.interview_library_parser as ilp
    lib_path = tmp_path / "interview_library.json"
    lib_path.write_text(json.dumps(data), encoding="utf-8")
    monkeypatch.setattr(ilp, "LIBRARY_PATH", str(lib_path))
    return lib_path


# ── Argparse ──────────────────────────────────────────────────────────────────

def test_argparse_requires_role():
    with pytest.raises(SystemExit):
        wc.build_parser().parse_args(["--stage", "hiring_manager"])


def test_argparse_requires_stage():
    with pytest.raises(SystemExit):
        wc.build_parser().parse_args(["--role", "TestRole"])


def test_argparse_valid_args():
    args = wc.build_parser().parse_args(["--role", "TestRole", "--stage", "hiring_manager"])
    assert args.role == "TestRole"
    assert args.stage == "hiring_manager"


# ── Docx location ─────────────────────────────────────────────────────────────

def test_locate_docx_exits_when_missing(tmp_path, monkeypatch):
    monkeypatch.setattr(wc, "JOBS_PACKAGES_DIR", str(tmp_path))
    with pytest.raises(SystemExit):
        wc._locate_docx("TestRole", "hiring_manager")


def test_locate_docx_returns_path_when_present(tmp_path, monkeypatch):
    pkg = tmp_path / "TestRole"
    pkg.mkdir()
    docx_path = pkg / "interview_prep_hiring_manager.docx"
    docx_path.write_bytes(b"")
    monkeypatch.setattr(wc, "JOBS_PACKAGES_DIR", str(tmp_path))
    result = wc._locate_docx("TestRole", "hiring_manager")
    assert result == str(docx_path)


# ── _extract_docx_paragraphs ──────────────────────────────────────────────────
# These tests use real (minimal) docx files built with python-docx.

def _make_minimal_docx(tmp_path, paragraphs):
    """Build a minimal .docx with the given (text, bold, italic) tuples."""
    from docx import Document
    doc = Document()
    for text, bold, italic in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return str(path)


def test_extract_skips_blank_paragraphs(tmp_path):
    path = _make_minimal_docx(tmp_path, [
        ("Hello", False, False),
        ("", False, False),
        ("World", False, False),
    ])
    result = wc._extract_docx_paragraphs(path)
    assert len(result) == 2
    assert result[0][0] == "Hello"
    assert result[1][0] == "World"


def test_extract_detects_italic_paragraphs(tmp_path):
    path = _make_minimal_docx(tmp_path, [
        ("Normal line", False, False),
        ("Italic coaching line", False, True),
    ])
    result = wc._extract_docx_paragraphs(path)
    assert result[0][2] is False   # not italic
    assert result[1][2] is True    # italic


# ── _split_sections ───────────────────────────────────────────────────────────

def _paras(texts, style="Normal"):
    """Helper: build paragraph tuples with given style and not italic."""
    return [(t, style, False) for t in texts]


def test_split_sections_finds_story_bank():
    paras = (
        _paras(["Interview Prep Package"], style="Heading 1") +
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 -- MBSE:"]) +
        _paras(["Situation: context here."]) +
        _paras(["Gap Preparation"], style="Heading 1") +
        _paras(["GAP 1 -- Networking [REQUIRED]:"])
    )
    sections = wc._split_sections(paras)
    assert any("STORY 1" in t for t, _, _ in sections["story_bank"])


def test_split_sections_finds_gap_prep():
    paras = (
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 -- MBSE:"]) +
        _paras(["Gap Preparation"], style="Heading 1") +
        _paras(["GAP 1 -- Networking [REQUIRED]:"]) +
        _paras(["Honest answer: here."])
    )
    sections = wc._split_sections(paras)
    assert any("GAP 1" in t for t, _, _ in sections["gap_prep"])


def test_split_sections_finds_questions():
    paras = (
        _paras(["Gap Preparation"], style="Heading 1") +
        _paras(["GAP 1 -- Topic [REQUIRED]:"]) +
        _paras(["Questions to Ask"], style="Heading 1") +
        _paras(["1. What does success look like at 6 months?"])
    )
    sections = wc._split_sections(paras)
    assert any("success" in t for t, _, _ in sections["questions"])


def test_split_sections_excludes_other_sections():
    paras = (
        _paras(["Company Role Brief"], style="Heading 1") +
        _paras(["Company overview here."]) +
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 -- MBSE:"])
    )
    sections = wc._split_sections(paras)
    assert not any("overview" in t for t, _, _ in sections["story_bank"])


def test_split_sections_content_with_salary_word_not_false_positive():
    paras = (
        _paras(["Story Bank"], style="Heading 1") +
        _paras(["STORY 1 -- Background:"]) +
        _paras(["They asked about salary."]) +
        _paras(["I provided a market range."])
    )
    sections = wc._split_sections(paras)
    # Content paragraph mentioning "salary" should NOT reset the section
    assert any("salary" in t.lower() for t, _, _ in sections["story_bank"])


# ── _parse_stories ────────────────────────────────────────────────────────────

def _story_paras(lines):
    """Build paragraph tuples for story bank content."""
    return [(line, "Normal", False) for line in lines]


def test_parse_stories_extracts_employer_title_dates():
    paras = _story_paras([
        "STORY 1 -- MBSE Toolchain [Systems Engineering]:",
        "Employer: G2 OPS | Systems Engineer | 2021-2023",
        "Situation: Context here.",
        "Task: What needed doing.",
        "Action: What I did.",
        "Result: The outcome.",
    ])
    stories = wc._parse_stories(paras)
    assert len(stories) == 1
    assert stories[0]["employer"] == "G2 OPS"
    assert stories[0]["title_held"] == "Systems Engineer"
    assert stories[0]["dates"] == "2021-2023"


def test_parse_stories_extracts_star_components():
    paras = _story_paras([
        "STORY 1 -- Requirement:",
        "Employer: ACME | Engineer | 2022-2023",
        "Situation: Sat text.",
        "Task: Task text.",
        "Action: Action text.",
        "Result: Result text.",
    ])
    s = wc._parse_stories(paras)[0]
    assert s["situation"] == "Sat text."
    assert s["task"] == "Task text."
    assert s["action"] == "Action text."
    assert s["result"] == "Result text."


def test_parse_stories_extracts_if_probed():
    paras = _story_paras([
        "STORY 1 -- Requirement:",
        "Employer: ACME | Engineer | 2022",
        "Situation: S.", "Task: T.", "Action: A.", "Result: R.",
        "If probed: One more sentence.",
    ])
    s = wc._parse_stories(paras)[0]
    assert s["if_probed"] == "One more sentence."


def test_parse_stories_if_probed_none_when_absent():
    paras = _story_paras([
        "STORY 1 -- Requirement:",
        "Employer: ACME | Engineer | 2022",
        "Situation: S.", "Task: T.", "Action: A.", "Result: R.",
    ])
    s = wc._parse_stories(paras)[0]
    assert s["if_probed"] is None


def test_parse_stories_skips_italic_paragraphs():
    paras = [
        ("STORY 1 -- Requirement:", "Normal", False),
        ("Employer: ACME | Engineer | 2022", "Normal", False),
        ("Signals: do not include this line.", "Normal", True),   # italic
        ("Situation: Real situation.", "Normal", False),
        ("Task: T.", "Normal", False),
        ("Action: A.", "Normal", False),
        ("Result: R.", "Normal", False),
    ]
    s = wc._parse_stories(paras)[0]
    assert s["situation"] == "Real situation."


def test_parse_stories_handles_multiple_stories():
    paras = _story_paras([
        "STORY 1 -- First:",
        "Employer: ACME | Eng | 2021",
        "Situation: S1.", "Task: T1.", "Action: A1.", "Result: R1.",
        "STORY 2 -- Second:",
        "Employer: Corp | Arch | 2022",
        "Situation: S2.", "Task: T2.", "Action: A2.", "Result: R2.",
    ])
    stories = wc._parse_stories(paras)
    assert len(stories) == 2
    assert stories[0]["employer"] == "ACME"
    assert stories[1]["employer"] == "Corp"


def test_parse_stories_continuation_after_empty_label():
    paras = _story_paras([
        "STORY 1 -- Requirement:",
        "Employer: ACME | Engineer | 2022",
        "Situation: S.", "Task: T.", "Action: A.", "Result: R.",
        "If probed:",
        "One more sentence.",
    ])
    s = wc._parse_stories(paras)[0]
    assert s["if_probed"] == "One more sentence."


# ── _parse_gaps ───────────────────────────────────────────────────────────────

def _gap_paras(lines):
    return [(line, "Normal", False) for line in lines]


def test_parse_gaps_extracts_label_and_severity_required():
    paras = _gap_paras([
        "GAP 1 -- IP Networking Expertise [REQUIRED]:",
        "Gap: I have not worked in IP networking.",
        "Honest answer: Direct acknowledgment.",
        "Bridge: Adjacent experience.",
        "Redirect: Strong redirect.",
    ])
    gaps = wc._parse_gaps(paras)
    assert len(gaps) == 1
    assert gaps[0]["gap_label"] == "IP Networking Expertise"
    assert gaps[0]["severity"] == "required"


def test_parse_gaps_extracts_severity_preferred():
    paras = _gap_paras([
        "GAP 1 -- Cameo TWC Experience [PREFERRED]:",
        "Gap: gap text.", "Honest answer: A.", "Bridge: B.", "Redirect: R.",
    ])
    assert wc._parse_gaps(paras)[0]["severity"] == "preferred"


def test_parse_gaps_extracts_triad():
    paras = _gap_paras([
        "GAP 1 -- Topic [REQUIRED]:",
        "Gap: gap text.",
        "Honest answer: Honest text.",
        "Bridge: Bridge text.",
        "Redirect: Redirect text.",
    ])
    g = wc._parse_gaps(paras)[0]
    assert g["honest_answer"] == "Honest text."
    assert g["bridge"] == "Bridge text."
    assert g["redirect"] == "Redirect text."


def test_parse_gaps_excludes_short_tenure_section():
    paras = _gap_paras([
        "SHORT TENURE EXPLANATION:",
        "I left after 8 months because...",
        "GAP 1 -- Networking [REQUIRED]:",
        "Gap: gap text.", "Honest answer: A.", "Bridge: B.", "Redirect: R.",
    ])
    gaps = wc._parse_gaps(paras)
    assert len(gaps) == 1
    assert gaps[0]["gap_label"] == "Networking"


def test_parse_gaps_stops_at_hard_questions():
    paras = _gap_paras([
        "GAP 1 -- Topic [REQUIRED]:",
        "Gap: g.", "Honest answer: A.", "Bridge: B.", "Redirect: R.",
        "HARD QUESTIONS TO PREPARE FOR:",
        "1. Tough question?",
    ])
    gaps = wc._parse_gaps(paras)
    assert len(gaps) == 1


def test_parse_gaps_skips_italic_paragraphs():
    paras = [
        ("GAP 1 -- Topic [REQUIRED]:", "Normal", False),
        ("Coaching note -- do not use.", "Normal", True),
        ("Gap: gap text.", "Normal", False),
        ("Honest answer: A.", "Normal", False),
        ("Bridge: B.", "Normal", False),
        ("Redirect: R.", "Normal", False),
    ]
    gaps = wc._parse_gaps(paras)
    assert len(gaps) == 1


# ── _parse_questions ──────────────────────────────────────────────────────────

def test_parse_questions_extracts_question_text():
    paras = [(
        "1. What does success look like at 6 months? Signals you're thinking about impact.",
        "Normal", False
    )]
    questions = wc._parse_questions(paras, "hiring_manager")
    assert len(questions) == 1
    assert questions[0]["text"] == "What does success look like at 6 months?"
    assert questions[0]["stage"] == "hiring_manager"


def test_parse_questions_strips_rationale():
    paras = [(
        "2. Where are the hard interface problems right now? This signals peer credibility.",
        "Normal", False
    )]
    q = wc._parse_questions(paras, "team_panel")[0]
    assert q["text"].endswith("?")
    assert "credibility" not in q["text"]


def test_parse_questions_skips_italic():
    paras = [
        ("1. Real question?", "Normal", False),
        ("Coaching note.", "Normal", True),
        ("2. Another question?", "Normal", False),
    ]
    questions = wc._parse_questions(paras, "recruiter")
    assert len(questions) == 2


def test_parse_questions_returns_empty_when_no_numbered_items():
    paras = [("Not a numbered item.", "Normal", False)]
    assert wc._parse_questions(paras, "hiring_manager") == []


def test_closing_question_excluded():
    paras = [
        ("1. What does success look like at 6 months?", "Normal", False),
        ("Based on our conversation, is there anything that gives you pause about my fit?",
         "Normal", False),
    ]
    questions = wc._parse_questions(paras, "hiring_manager")
    assert len(questions) == 1
    assert not any("Based on our conversation" in q["text"] for q in questions)
