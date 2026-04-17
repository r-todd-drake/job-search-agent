# tests/phase5/test_interview_prep.py

import json
import pytest
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch
from docx import Document

FIXTURE_JD = Path(__file__).parent.parent / "fixtures" / "stage_files" / "sample_jd.txt"
FIXTURE_STAGE2 = Path(__file__).parent.parent / "fixtures" / "stage_files" / "stage2_approved.txt"
FIXTURE_LIBRARY = Path(__file__).parent.parent / "fixtures" / "library" / "experience_library.json"
FIXTURE_PROFILE = Path(__file__).parent.parent / "fixtures" / "library" / "candidate_profile.md"

MOCK_PREP_RESPONSE = """## SECTION 1: COMPANY & ROLE BRIEF
Acme Defense Systems is a defense contractor focused on autonomous maritime systems.
Role: Senior Systems Engineer supporting MBSE and DoDAF development.
Salary guidance: $150,000 - $180,000.

## SECTION 1.5: INTRODUCE YOURSELF
I am a senior systems engineer with 20 years of defense SE experience.
I specialize in MBSE and DoDAF architectural development.

## SECTION 2: STORY BANK
STAR Story 1 - MBSE Leadership
Situation: Led MBSE development for autonomous surface vessel program.
Task: Develop DoDAF architectural views using Cameo Systems Modeler.
Action: Facilitated IPT working groups with government stakeholders.
Result: Delivered system-of-systems architecture supporting multi-domain C2 integration.

STAR Story 2 - Stakeholder Engagement
Situation: Government stakeholder alignment required for requirements definition.
Task: Define operational requirements and ConOps.
Action: Conducted workshops and facilitated reviews.
Result: Approved ConOps baseline.

STAR Story 3 - Architecture Integration
Situation: System integration complexity.
Task: Develop integration architecture.
Action: Applied DoDAF SV views.
Result: Successful integration milestone.

## SECTION 3: GAP PREPARATION
REQUIRED: All required qualifications met.
PREFERRED: JADC2 experience limited -- acknowledge and reframe.

## SECTION 4: QUESTIONS TO ASK
1. What is the acquisition phase for this program?
2. How is MBSE integrated into the program baseline?
"""


def make_mock_client(response_text):
    client = MagicMock()
    client.messages.create.return_value = MagicMock(
        content=[MagicMock(text=response_text)]
    )
    return client


def make_role_data():
    return {
        "jd_text": FIXTURE_JD.read_text(encoding="utf-8"),
        "stage_text": FIXTURE_STAGE2.read_text(encoding="utf-8"),
        "library": json.loads(FIXTURE_LIBRARY.read_text(encoding="utf-8")),
        "candidate_profile": FIXTURE_PROFILE.read_text(encoding="utf-8"),
        "role_name": "acme_sse",
    }


def test_generate_prep_creates_both_output_files():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))
        assert txt_path.exists()
        assert docx_path.exists()


def test_generate_prep_txt_has_required_sections():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))
        content = txt_path.read_text(encoding="utf-8")

    assert "COMPANY" in content or "BRIEF" in content
    assert "STORY" in content or "STAR" in content
    assert "GAP" in content
    assert "QUESTION" in content


def test_generate_prep_star_stories_reference_resume_content():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))
        content = txt_path.read_text(encoding="utf-8")

    assert "Cameo Systems Modeler" in content or "MBSE" in content


def test_generate_prep_no_pii_in_api_payload(pii_values, monkeypatch):
    import importlib
    import scripts.utils.pii_filter as pii_module
    for key, val in [
        ("CANDIDATE_NAME", pii_values["name"]),
        ("CANDIDATE_PHONE", pii_values["phone"]),
        ("CANDIDATE_EMAIL", pii_values["email"]),
        ("CANDIDATE_LINKEDIN", pii_values["linkedin"]),
        ("CANDIDATE_GITHUB", pii_values["github"]),
    ]:
        monkeypatch.setenv(key, val)
    importlib.reload(pii_module)

    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    role_data["stage_text"] = (
        f"Contact: {pii_values['name']} | {pii_values['email']}\n"
        + role_data["stage_text"]
    )

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))

    call_args = client.messages.create.call_args
    full_payload = str(call_args)
    for pii_value in pii_values.values():
        assert pii_value not in full_payload, f"PII found in payload: {pii_value}"


def test_generate_prep_docx_readable():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))
        doc = Document(str(docx_path))
        assert len(doc.paragraphs) > 0


def test_dry_run_no_api_calls():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager",
                      str(txt_path), str(docx_path), dry_run=True)
        assert not txt_path.exists(), "dry_run must not write output files"

    assert client.messages.create.call_count == 0, \
        f"dry_run must make no API calls, got {client.messages.create.call_count}"


def test_no_module_level_execution_on_import():
    import scripts.phase5_interview_prep  # noqa: F401


def test_stage_profile_has_required_keys():
    from scripts.phase5_interview_prep import STAGE_PROFILES, VALID_STAGES

    required_keys = {
        "label", "description", "story_count", "story_depth",
        "gap_behavior", "salary_in_section1", "section1_focus", "questions_prompt",
    }
    assert set(VALID_STAGES) == {"recruiter", "hiring_manager", "team_panel"}
    for stage, profile in STAGE_PROFILES.items():
        missing = required_keys - profile.keys()
        assert not missing, f"Stage '{stage}' missing keys: {missing}"
    assert "peer_frame_prompt" in STAGE_PROFILES["team_panel"]
    assert STAGE_PROFILES["recruiter"]["gap_behavior"] == "omit"
    assert STAGE_PROFILES["hiring_manager"]["salary_in_section1"] is True
    assert STAGE_PROFILES["team_panel"]["story_depth"] == "full_technical"


def test_invalid_stage_raises_system_exit(monkeypatch):
    import scripts.phase5_interview_prep as mod
    monkeypatch.setattr("sys.argv", ["phase5_interview_prep.py", "--role", "test_role",
                                      "--interview_stage", "badstage"])
    with pytest.raises(SystemExit):
        mod.main()


def test_stage_specific_filenames():
    from scripts.phase5_interview_prep import _output_paths
    for stage in ("recruiter", "hiring_manager", "team_panel"):
        txt, docx = _output_paths("/some/dir", stage)
        assert txt.endswith(f"interview_prep_{stage}.txt"), \
            f"Expected interview_prep_{stage}.txt, got {txt}"
        assert docx.endswith(f"interview_prep_{stage}.docx"), \
            f"Expected interview_prep_{stage}.docx, got {docx}"


def test_stage_in_output_header():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_recruiter.txt"
        docx_path = Path(tmpdir) / "interview_prep_recruiter.docx"
        generate_prep(client, role_data, "recruiter", str(txt_path), str(docx_path))
        content = txt_path.read_text(encoding="utf-8")

    assert "Stage: Recruiter Screen" in content
    assert "Short screen" in content


def test_extract_profile_section_found():
    from scripts.phase5_interview_prep import extract_profile_section
    text = "## INTRO MONOLOGUE\nHello world.\n## OTHER SECTION\nOther content."
    result = extract_profile_section(text, "INTRO MONOLOGUE")
    assert "Hello world" in result
    assert "OTHER SECTION" not in result


def test_extract_profile_section_missing():
    from scripts.phase5_interview_prep import extract_profile_section
    result = extract_profile_section("## OTHER SECTION\nStuff.", "INTRO MONOLOGUE")
    assert result == ""


def test_extract_profile_section_last_section():
    from scripts.phase5_interview_prep import extract_profile_section
    text = "## SHORT TENURE EXPLANATION\nI left because the contract ended."
    result = extract_profile_section(text, "SHORT TENURE EXPLANATION")
    assert "contract ended" in result


def test_section1_salary_only_for_hiring_manager():
    from scripts.phase5_interview_prep import generate_prep

    # hiring_manager -- salary guidance should appear in Section 1 API call
    client_hm = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    with tempfile.TemporaryDirectory() as tmpdir:
        generate_prep(client_hm, role_data, "hiring_manager",
                      str(Path(tmpdir) / "interview_prep_hiring_manager.txt"),
                      str(Path(tmpdir) / "interview_prep_hiring_manager.docx"))

    hm_calls = client_hm.messages.create.call_args_list
    section1_call = str(hm_calls[0])
    assert "SALARY" in section1_call.upper()

    # recruiter -- salary guidance should NOT be in Section 1 API call
    client_rec = make_mock_client(MOCK_PREP_RESPONSE)
    with tempfile.TemporaryDirectory() as tmpdir:
        generate_prep(client_rec, role_data, "recruiter",
                      str(Path(tmpdir) / "interview_prep_recruiter.txt"),
                      str(Path(tmpdir) / "interview_prep_recruiter.docx"))

    rec_calls = client_rec.messages.create.call_args_list
    section1_rec_call = str(rec_calls[0])
    assert "SALARY EXPECTATIONS GUIDANCE" not in section1_rec_call


def test_intro_monologue_in_output():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))
        content = txt_path.read_text(encoding="utf-8")

    assert "INTRODUCE YOURSELF" in content or "SECTION 1.5" in content


def test_section2_story_count_in_api_payload():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        generate_prep(client, role_data, "recruiter",
                      str(Path(tmpdir) / "interview_prep_recruiter.txt"),
                      str(Path(tmpdir) / "interview_prep_recruiter.docx"))

    # Call order for recruiter: S1=0, S1.5=1, S2=2
    section2_call = str(client.messages.create.call_args_list[2])
    assert "1-2" in section2_call
    assert "headline" in section2_call.lower()
    assert "Do NOT reference gaps" in section2_call or "Suppress gap" in section2_call or "NOT reference gaps" in section2_call


@pytest.mark.live
def test_generate_prep_live():
    """Tier 2: real API call with web search."""
    import os
    from anthropic import Anthropic
    from scripts.phase5_interview_prep import generate_prep

    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    role_data = make_role_data()
    role_data["role_name"] = "acme_sse_test"

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        docx_path = Path(tmpdir) / "interview_prep_hiring_manager.docx"
        generate_prep(client, role_data, "hiring_manager", str(txt_path), str(docx_path))
        content = txt_path.read_text(encoding="utf-8")

    assert len(content) > 500
    assert "STAR" in content or "Story" in content
    assert "GAP" in content or "Gap" in content


def test_recruiter_skips_gap_api_call():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_recruiter.txt"
        generate_prep(client, role_data, "recruiter",
                      str(txt_path),
                      str(Path(tmpdir) / "interview_prep_recruiter.docx"))
        content = txt_path.read_text(encoding="utf-8")

    # recruiter: S1=0, S1.5=1, S2=2, S3 SKIPPED, S4=3 -- total 4 calls
    assert client.messages.create.call_count == 4, (
        f"Expected 4 API calls for recruiter, got {client.messages.create.call_count}"
    )
    assert "do not volunteer gaps" in content.lower()


def test_short_tenure_block_in_output():
    from scripts.phase5_interview_prep import generate_prep

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_hiring_manager.txt"
        generate_prep(client, role_data, "hiring_manager",
                      str(txt_path),
                      str(Path(tmpdir) / "interview_prep_hiring_manager.docx"))
        content = txt_path.read_text(encoding="utf-8")

    # The fixture profile has ## SHORT TENURE EXPLANATION
    assert "SHORT TENURE EXPLANATION" in content


def test_team_panel_peer_frame_bold_in_docx():
    from scripts.phase5_interview_prep import generate_prep

    # Mock response that contains a Peer Frame label
    mock_with_peer_frame = MOCK_PREP_RESPONSE + "\nPeer Frame: I understand this gap operationally.\n"
    client = make_mock_client(mock_with_peer_frame)
    role_data = make_role_data()

    with tempfile.TemporaryDirectory() as tmpdir:
        txt_path = Path(tmpdir) / "interview_prep_team_panel.txt"
        docx_path = Path(tmpdir) / "interview_prep_team_panel.docx"
        generate_prep(client, role_data, "team_panel", str(txt_path), str(docx_path))
        doc = Document(str(docx_path))

    # Find paragraphs that start with "Peer Frame:"
    peer_frame_paragraphs = [
        p for p in doc.paragraphs if p.text.startswith("Peer Frame:")
    ]
    assert peer_frame_paragraphs, "No 'Peer Frame:' paragraph found in docx"
    assert peer_frame_paragraphs[0].runs[0].bold, "Peer Frame: label run should be bold"


# ---- _extract_jd_tags tests ----

def test_extract_jd_tags_returns_matching_tags():
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import _extract_jd_tags
    with patch.object(ilp, "load_tags", return_value=["mbse", "systems-engineering", "leadership"]):
        result = _extract_jd_tags("This role requires MBSE experience and Systems Engineering skills.")
    assert "mbse" in result
    assert "systems-engineering" in result
    assert "leadership" not in result


def test_extract_jd_tags_returns_empty_when_no_match():
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import _extract_jd_tags
    with patch.object(ilp, "load_tags", return_value=["mbse", "clearance"]):
        result = _extract_jd_tags("General software engineering role.")
    assert result == []


def test_extract_jd_tags_case_insensitive():
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import _extract_jd_tags
    with patch.object(ilp, "load_tags", return_value=["mbse"]):
        result = _extract_jd_tags("Experience with MBSE frameworks required.")
    assert "mbse" in result


# ---- library seeding in _build_section2_prompt ----

def test_build_section2_prompt_injects_seed_block():
    from scripts.phase5_interview_prep import _build_section2_prompt, STAGE_PROFILES
    seeds = [{
        "id": "story_001", "employer": "G2 OPS", "title": "Lead SE", "dates": "2022-2024",
        "situation": "Led MBSE.", "task": "Build OV-1.", "action": "Facilitated IPT.",
        "result": "Delivered baseline.", "tags": ["mbse"], "if_probed": None
    }]
    prompt = _build_section2_prompt("JD text", "story context", "profile",
                                    STAGE_PROFILES["hiring_manager"], library_seeds=seeds)
    assert "VETTED LIBRARY STORIES" in prompt
    assert "G2 OPS" in prompt
    assert "library-seeded" in prompt


def test_build_section2_prompt_no_seed_block_when_none():
    from scripts.phase5_interview_prep import _build_section2_prompt, STAGE_PROFILES
    prompt = _build_section2_prompt("JD text", "story context", "profile",
                                    STAGE_PROFILES["hiring_manager"], library_seeds=None)
    assert "VETTED LIBRARY STORIES" not in prompt


# ---- library seeding in _build_gap_prompt ----

def test_build_gap_prompt_injects_seed_block():
    from scripts.phase5_interview_prep import _build_gap_prompt, STAGE_PROFILES
    seeds = [{
        "id": "gap_001", "gap_label": "no SCIF experience", "severity": "REQUIRED",
        "honest_answer": "I have not worked in SCIF.", "bridge": "Worked TS cleared.",
        "redirect": "Adapt quickly.", "tags": ["clearance"]
    }]
    prompt = _build_gap_prompt("JD text", "gaps text", "profile",
                               STAGE_PROFILES["hiring_manager"], library_seeds=seeds)
    assert "VETTED GAP RESPONSES" in prompt
    assert "no SCIF experience" in prompt
    assert "library-seeded" in prompt


def test_build_gap_prompt_no_seed_block_when_none():
    from scripts.phase5_interview_prep import _build_gap_prompt, STAGE_PROFILES
    prompt = _build_gap_prompt("JD text", "gaps text", "profile",
                               STAGE_PROFILES["hiring_manager"], library_seeds=None)
    assert "VETTED GAP RESPONSES" not in prompt


# ---- generate_prep calls get_stories with jd tags ----

def test_generate_prep_calls_library_with_jd_tags(monkeypatch):
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    captured_tags = {}
    def fake_get_stories(tags=None, **kw):
        captured_tags["tags"] = tags
        return []
    monkeypatch.setattr(ilp, "load_tags", lambda: ["mbse", "systems-engineering"])
    monkeypatch.setattr(ilp, "get_stories", fake_get_stories)
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    with tempfile.TemporaryDirectory() as tmpdir:
        generate_prep(client, role_data, "hiring_manager",
                      str(Path(tmpdir) / "p.txt"), str(Path(tmpdir) / "p.docx"))

    assert captured_tags.get("tags") is not None


# ---- cold path: no seeds, prompt unchanged behaviour ----

def test_generate_prep_cold_path_no_library_seeds(monkeypatch):
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    with tempfile.TemporaryDirectory() as tmpdir:
        txt = Path(tmpdir) / "p.txt"
        generate_prep(client, role_data, "hiring_manager",
                      str(txt), str(Path(tmpdir) / "p.docx"))
    # No library seeds -- section 2 call should NOT include VETTED LIBRARY STORIES
    section2_kwargs = client.messages.create.call_args_list[2].kwargs
    prompt = section2_kwargs["messages"][0]["content"]
    assert "VETTED LIBRARY STORIES" not in prompt


# ---- performance signal injected into story seed ----

def test_performance_signal_injected_into_story_seed_prompt(monkeypatch, tmp_path):
    import scripts.interview_library_parser as ilp
    import scripts.phase5_debrief_utils as dbu
    from scripts.phase5_interview_prep import generate_prep

    sample_story = {
        "id": "story_001", "employer": "G2 OPS", "title": "Lead SE", "dates": "2022-2024",
        "situation": "Led MBSE.", "task": "OV-1.", "action": "IPT.", "result": "Baseline.",
        "tags": ["mbse"], "if_probed": None, "roles_used": []
    }
    monkeypatch.setattr(ilp, "load_tags", lambda: ["mbse"])
    monkeypatch.setattr(ilp, "get_stories", lambda tags=None, **kw: [sample_story] if tags else [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs",
                        lambda: [{"metadata": {"role": "x", "stage": "hm", "interview_date": "2026-04-01",
                                               "panel_label": None, "interviewers": []},
                                  "stories_used": [{"library_id": "story_001", "landed": "yes", "tags": []}],
                                  "gaps_surfaced": [], "salary_exchange": {},
                                  "advancement_read": {}, "what_i_said": None, "open_notes": None}])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    generate_prep(client, role_data, "hiring_manager",
                  str(tmp_path / "p.txt"), str(tmp_path / "p.docx"))

    section2_kwargs = client.messages.create.call_args_list[2].kwargs
    prompt = section2_kwargs["messages"][0]["content"]
    assert "Used 1 times" in prompt
    assert "yes x1" in prompt


# ---- salary actuals override ----

def test_salary_actuals_override_injects_actuals_into_section1_prompt(monkeypatch, tmp_path):
    import scripts.phase5_debrief_utils as dbu
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    actuals_debrief = {
        "metadata": {"role": "acme_sse", "stage": "recruiter", "interview_date": "2026-04-10",
                     "panel_label": None, "interviewers": []},
        "advancement_read": {}, "stories_used": [], "gaps_surfaced": [],
        "salary_exchange": {"range_given_min": 145000, "range_given_max": 175000,
                            "candidate_anchor": 168000, "candidate_floor": 152000},
        "what_i_said": None, "open_notes": None,
    }
    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [actuals_debrief])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    generate_prep(client, role_data, "hiring_manager",
                  str(tmp_path / "p.txt"), str(tmp_path / "p.docx"))

    section1_kwargs = client.messages.create.call_args_list[0].kwargs
    prompt = section1_kwargs["messages"][0]["content"]
    assert "SALARY ACTUALS" in prompt
    assert "145,000" in prompt
    assert "175,000" in prompt


def test_no_salary_override_when_no_debrief(monkeypatch, tmp_path):
    import scripts.phase5_debrief_utils as dbu
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    generate_prep(client, role_data, "hiring_manager",
                  str(tmp_path / "p.txt"), str(tmp_path / "p.docx"))

    section1_kwargs = client.messages.create.call_args_list[0].kwargs
    prompt = section1_kwargs["messages"][0]["content"]
    assert "SALARY ACTUALS" not in prompt


def test_no_performance_signal_when_no_debrief_history(monkeypatch, tmp_path):
    import scripts.interview_library_parser as ilp
    import scripts.phase5_debrief_utils as dbu
    from scripts.phase5_interview_prep import generate_prep

    sample_story = {
        "id": "story_999", "employer": "Acme", "title": "SE", "dates": "2020-2022",
        "situation": "s", "task": "t", "action": "a", "result": "r",
        "tags": ["mbse"], "if_probed": None, "roles_used": []
    }
    monkeypatch.setattr(ilp, "load_tags", lambda: ["mbse"])
    monkeypatch.setattr(ilp, "get_stories", lambda tags=None, **kw: [sample_story] if tags else [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    generate_prep(client, role_data, "hiring_manager",
                  str(tmp_path / "p.txt"), str(tmp_path / "p.docx"))

    section2_kwargs = client.messages.create.call_args_list[2].kwargs
    prompt = section2_kwargs["messages"][0]["content"]
    assert "Performance:" not in prompt


# ---- continuity section in output ----

def test_continuity_section_appears_in_txt_when_debriefs_exist(monkeypatch, tmp_path):
    import scripts.phase5_debrief_utils as dbu
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    prior_debrief = {
        "metadata": {"role": "acme_sse", "stage": "recruiter", "interview_date": "2026-04-01",
                     "panel_label": None,
                     "interviewers": [{"name": "HR Alice", "title": "Recruiter", "notes": ""}]},
        "advancement_read": {"assessment": "maybe", "notes": ""},
        "stories_used": [], "gaps_surfaced": [], "salary_exchange": {},
        "what_i_said": "Said I want hybrid work.",
        "open_notes": None,
    }
    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [prior_debrief])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    txt_path = tmp_path / "p.txt"
    generate_prep(client, role_data, "hiring_manager",
                  str(txt_path), str(tmp_path / "p.docx"))
    content = txt_path.read_text(encoding="utf-8")
    assert "CONTINUITY SUMMARY" in content
    assert "Said I want hybrid work." in content
    assert "HR Alice" in content


def test_no_continuity_section_when_no_debriefs(monkeypatch, tmp_path):
    import scripts.phase5_debrief_utils as dbu
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    txt_path = tmp_path / "p.txt"
    generate_prep(client, role_data, "hiring_manager",
                  str(txt_path), str(tmp_path / "p.docx"))
    content = txt_path.read_text(encoding="utf-8")
    assert "CONTINUITY SUMMARY" not in content


def test_continuity_section_in_docx(monkeypatch, tmp_path):
    import scripts.phase5_debrief_utils as dbu
    import scripts.interview_library_parser as ilp
    from scripts.phase5_interview_prep import generate_prep

    prior_debrief = {
        "metadata": {"role": "acme_sse", "stage": "recruiter", "interview_date": "2026-04-01",
                     "panel_label": None,
                     "interviewers": [{"name": "HR Alice", "title": "Recruiter", "notes": ""}]},
        "advancement_read": {}, "stories_used": [], "gaps_surfaced": [], "salary_exchange": {},
        "what_i_said": "Said I prefer remote.", "open_notes": None,
    }
    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [prior_debrief])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    docx_path = tmp_path / "p.docx"
    generate_prep(client, role_data, "hiring_manager",
                  str(tmp_path / "p.txt"), str(docx_path))
    doc = Document(str(docx_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Continuity Summary" in all_text or "CONTINUITY" in all_text


# ---- terminal notifications (debrief_utils logic validation) ----

def test_find_unmatched_returns_ghost_story_id():
    import scripts.phase5_debrief_utils as dbu
    unmatched_debrief = {
        "metadata": {"role": "r", "stage": "hiring_manager", "interview_date": "2026-04-10",
                     "panel_label": None, "interviewers": []},
        "stories_used": [{"library_id": "ghost_id", "tags": ["mbse"], "landed": "yes"}],
        "gaps_surfaced": [], "salary_exchange": {}, "advancement_read": {},
        "what_i_said": None, "open_notes": None,
    }
    with patch("scripts.interview_library_parser._load_library",
               return_value={"stories": [], "gap_responses": [], "questions": []}):
        stories, gaps = dbu.find_unmatched_debrief_content([unmatched_debrief])
    assert len(stories) == 1
    assert stories[0]["library_id"] == "ghost_id"


def test_thankyou_notification_check():
    import scripts.phase5_debrief_utils as dbu
    d = {
        "metadata": {"stage": "hiring_manager", "panel_label": None, "interview_date": "2026-04-10",
                     "interviewers": [], "role": "r"},
        "stories_used": [], "gaps_surfaced": [], "salary_exchange": {}, "advancement_read": {},
        "what_i_said": None, "open_notes": None,
    }
    assert dbu.has_debrief_for_stage([d], "hiring_manager") is True
    assert dbu.has_debrief_for_stage([d], "team_panel") is False


# ---- no-regression: absent library, absent debriefs ----

def test_no_regression_absent_library_file(monkeypatch, tmp_path):
    import scripts.interview_library_parser as ilp
    import scripts.phase5_debrief_utils as dbu
    from scripts.phase5_interview_prep import generate_prep

    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    txt = tmp_path / "p.txt"
    docx = tmp_path / "p.docx"
    generate_prep(client, role_data, "hiring_manager", str(txt), str(docx))
    assert txt.exists()
    assert docx.exists()


def test_no_regression_absent_debriefs_dir(monkeypatch, tmp_path):
    import scripts.interview_library_parser as ilp
    import scripts.phase5_debrief_utils as dbu
    from scripts.phase5_interview_prep import generate_prep

    monkeypatch.setattr(ilp, "load_tags", lambda: [])
    monkeypatch.setattr(ilp, "get_stories", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_gap_responses", lambda **kw: [])
    monkeypatch.setattr(ilp, "get_questions", lambda **kw: [])
    monkeypatch.setattr(dbu, "load_all_debriefs", lambda: [])
    monkeypatch.setattr(dbu, "load_debriefs", lambda role: [])

    client = make_mock_client(MOCK_PREP_RESPONSE)
    role_data = make_role_data()
    txt = tmp_path / "p.txt"
    generate_prep(client, role_data, "recruiter", str(txt), str(tmp_path / "p.docx"))
    assert txt.exists()
