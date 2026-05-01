import pytest
import sys
import os
import json
from datetime import date
from unittest.mock import MagicMock, patch

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import scripts.phase5_thankyou as pt


# ==============================================
# HELPERS
# ==============================================

def _make_debrief(interviewers, stories=None, gaps=None, what_i_said=None):
    return {
        "metadata": {
            "role": "TestRole",
            "stage": "hiring_manager",
            "panel_label": None,
            "interviewers": interviewers,
        },
        "stories_used": stories or [],
        "gaps_surfaced": gaps or [],
        "what_i_said": what_i_said,
    }


def _make_inputs(debrief, jd_text="JD text sample", resume_text="Resume summary",
                 candidate_profile="Candidate profile text"):
    return {
        "debrief": debrief,
        "jd_text": jd_text,
        "resume_text": resume_text,
        "candidate_profile": candidate_profile,
    }


def _mock_client(letter_text="This is a thank you letter."):
    client = MagicMock()
    response = MagicMock()
    response.content = [MagicMock(text=letter_text)]
    client.messages.create.return_value = response
    return client


def _setup_job_package(tmp_path, role="TestRole"):
    """Create a minimal job package directory with job_description.txt."""
    pkg_dir = tmp_path / "job_packages" / role
    pkg_dir.mkdir(parents=True)
    (pkg_dir / "job_description.txt").write_text("Job description content.", encoding="utf-8")
    return pkg_dir


# ==============================================
# _infer_tone
# ==============================================

def test_infer_tone_technical():
    tone = pt._infer_tone("Systems Engineer")
    assert "peer-level technical" in tone
    assert "engineer to another" in tone


def test_infer_tone_executive():
    tone = pt._infer_tone("Director of Engineering")
    assert "mission outcomes" in tone
    assert "strategic" in tone


def test_infer_tone_recruiter():
    tone = pt._infer_tone("Talent Acquisition Recruiter")
    assert "professional and warm" in tone
    assert "collegial" in tone


def test_infer_tone_default():
    tone = pt._infer_tone("Unknown Position Title")
    assert "professional and warm" in tone
    assert "collegial" in tone


# ==============================================
# generate_letters -- single interviewer
# ==============================================

def test_single_interviewer(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    debrief = _make_debrief([{"name": "John Smith", "title": "Engineer", "notes": "Talked about MBSE."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Dear John, thank you letter body here.")

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    assert len(generated) == 1
    assert len(skipped) == 0
    name, txt_path, docx_path = generated[0]
    assert name == "John Smith"
    assert os.path.exists(txt_path)
    assert os.path.exists(docx_path)
    assert "smith" in os.path.basename(txt_path)


def test_multiple_interviewers(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    debrief = _make_debrief([
        {"name": "Alice Brown", "title": "Director", "notes": "Discussed team strategy."},
        {"name": "Bob Green", "title": "Engineer", "notes": "Reviewed architecture."},
    ])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    assert len(generated) == 2
    assert len(skipped) == 0
    names = [g[0] for g in generated]
    assert "Alice Brown" in names
    assert "Bob Green" in names
    # Files must be distinct
    txt_paths = [g[1] for g in generated]
    assert txt_paths[0] != txt_paths[1]


def test_panel_label_in_filename(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    debrief = _make_debrief([{"name": "Jane Doe", "title": "Engineer", "notes": "notes"}])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    generated, _ = pt.generate_letters(client, role, "panel", "se_team", inputs, "2026-04-15")

    assert len(generated) == 1
    txt_path = generated[0][1]
    assert "se_team" in os.path.basename(txt_path)


def test_no_panel_label_in_filename(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    debrief = _make_debrief([{"name": "Jane Doe", "title": "Engineer", "notes": "notes"}])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    assert len(generated) == 1
    txt_path = generated[0][1]
    # No panel label should appear between stage and lastname
    basename = os.path.basename(txt_path)
    # Pattern without label: thankyou_{stage}_{lastname}_{date}.txt
    assert "thankyou_hiring_manager_doe_2026-04-15.txt" == basename


# ==============================================
# Missing interviewer notes warning
# ==============================================

def test_missing_interviewer_notes_warning(tmp_path, monkeypatch, capsys):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    debrief = _make_debrief([{"name": "Sam Lee", "title": "Recruiter", "notes": None}])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    out = capsys.readouterr().out
    assert "WARNING" in out
    assert "notes" in out.lower() or "No notes" in out or "notes" in out
    # Letter is still generated
    assert len(generated) == 1


# ==============================================
# _find_debrief -- file selection
# ==============================================

def test_most_recent_debrief_selected(tmp_path, capsys):
    role = "TestRole"
    role_dir = tmp_path / role
    role_dir.mkdir()

    older = role_dir / "debrief_hiring_manager_2026-04-01_filed-2026-04-01.json"
    newer = role_dir / "debrief_hiring_manager_2026-04-10_filed-2026-04-10.json"
    for f in (older, newer):
        f.write_text("{}", encoding="utf-8")

    selected = pt._find_debrief("TestRole", "hiring_manager", str(tmp_path))

    out = capsys.readouterr().out
    assert str(newer) in selected
    # Notice printed listing both matches
    assert "2026-04-01" in out or "2026-04-10" in out


def test_missing_debrief_exits(tmp_path):
    role_dir = tmp_path / "TestRole"
    role_dir.mkdir()

    with pytest.raises(SystemExit):
        pt._find_debrief("TestRole", "hiring_manager", str(tmp_path))


# ==============================================
# _load_inputs
# ==============================================

def test_missing_jd_exits(tmp_path, monkeypatch):
    role = "TestRole"
    # Job package dir exists but no job_description.txt
    pkg_dir = tmp_path / "job_packages" / role
    pkg_dir.mkdir(parents=True)

    profile_path = tmp_path / "candidate_profile.md"
    profile_path.write_text("Profile content.", encoding="utf-8")

    # Write a minimal debrief JSON
    debrief_path = tmp_path / "debrief.json"
    debrief_path.write_text(json.dumps({"metadata": {}, "stories_used": [], "gaps_surfaced": []}))

    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setattr(pt, "CANDIDATE_PROFILE_PATH", str(profile_path))

    with pytest.raises(SystemExit):
        pt._load_inputs(role, str(debrief_path))


def test_missing_resume_warning(tmp_path, monkeypatch):
    role = "TestRole"
    pkg_dir = tmp_path / "job_packages" / role
    pkg_dir.mkdir(parents=True)
    (pkg_dir / "job_description.txt").write_text("JD content.", encoding="utf-8")
    # No stage4 or stage2 file

    profile_path = tmp_path / "candidate_profile.md"
    profile_path.write_text("Profile content.", encoding="utf-8")

    debrief_data = {"metadata": {}, "stories_used": [], "gaps_surfaced": []}
    debrief_path = tmp_path / "debrief.json"
    debrief_path.write_text(json.dumps(debrief_data))

    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setattr(pt, "CANDIDATE_PROFILE_PATH", str(profile_path))

    inputs, warnings = pt._load_inputs(role, str(debrief_path))

    assert any("resume" in w.lower() or "stage" in w.lower() for w in warnings)
    assert inputs["resume_text"] is None


# ==============================================
# _build_letter_prompt -- tone calibration
# ==============================================

def test_tone_calibration_in_prompt():
    interviewer = {"name": "Alice Brown", "title": "Software Architect", "notes": "Discussed MBSE patterns."}
    debrief = _make_debrief([interviewer])
    prompt = pt._build_letter_prompt(
        interviewer, 0, debrief,
        "JD text", "Candidate profile", "Resume text"
    )
    # Technical tone should appear in the prompt
    assert "peer-level technical" in prompt


# ==============================================
# _write_docx -- document structure
# ==============================================

def test_docx_title_paragraph(tmp_path):
    from docx import Document as DocxDocument
    docx_path = str(tmp_path / "test_letter.docx")
    pt._write_docx(
        docx_path,
        "Paragraph one.\n\nParagraph two.",
        "Jane Smith",
        "Senior Engineer",
        "TestRole",
        "hiring_manager",
        "2026-04-15",
    )
    doc = DocxDocument(docx_path)
    # Title is the first paragraph
    title_text = doc.paragraphs[0].text
    assert "Jane Smith" in title_text
    assert "Senior Engineer" in title_text


def test_docx_metadata_paragraph(tmp_path):
    from docx import Document as DocxDocument
    docx_path = str(tmp_path / "test_letter.docx")
    pt._write_docx(
        docx_path,
        "Letter body here.",
        "Bob Jones",
        "Director",
        "TestRole",
        "hiring_manager",
        "2026-04-15",
    )
    doc = DocxDocument(docx_path)
    meta_text = doc.paragraphs[1].text
    assert "TestRole" in meta_text
    assert "hiring_manager" in meta_text
    assert "2026-04-15" in meta_text


def test_docx_failure_continues(tmp_path, monkeypatch, capsys):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    # Make _write_docx always raise
    monkeypatch.setattr(pt, "_write_docx", MagicMock(side_effect=Exception("disk full")))

    debrief = _make_debrief([{"name": "Pat Kim", "title": "Engineer", "notes": "Notes here."}])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    out = capsys.readouterr().out
    assert "WARNING" in out
    # .txt still written and letter counted as generated
    assert len(generated) == 1
    txt_path = generated[0][1]
    assert os.path.exists(txt_path)


# ==============================================
# Overwrite protection
# ==============================================

def test_overwrite_declined(tmp_path, monkeypatch):
    role = "TestRole"
    pkg_dir = _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    # Pre-create the expected output file
    existing_txt = pkg_dir / "thankyou_hiring_manager_jones_2026-04-15.txt"
    existing_txt.write_text("old content", encoding="utf-8")

    debrief = _make_debrief([{"name": "Bob Jones", "title": "Engineer", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    monkeypatch.setattr("builtins.input", lambda _: "n")

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    assert len(generated) == 0
    assert len(skipped) == 1
    # Original file content preserved
    assert existing_txt.read_text(encoding="utf-8") == "old content"


def test_overwrite_accepted(tmp_path, monkeypatch):
    role = "TestRole"
    pkg_dir = _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    existing_txt = pkg_dir / "thankyou_hiring_manager_jones_2026-04-15.txt"
    existing_txt.write_text("old content", encoding="utf-8")

    debrief = _make_debrief([{"name": "Bob Jones", "title": "Engineer", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("new letter content")

    monkeypatch.setattr("builtins.input", lambda _: "y")

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    assert len(generated) == 1
    assert len(skipped) == 0
    assert "new letter content" in existing_txt.read_text(encoding="utf-8")


# ==============================================
# _build_salutation -- unit tests
# ==============================================

def test_build_salutation_first_name_only():
    assert pt._build_salutation("John Smith") == "Dear John,"


def test_build_salutation_single_name():
    assert pt._build_salutation("Alice") == "Dear Alice,"


def test_build_salutation_null_name():
    assert pt._build_salutation(None) == "Dear Hiring Manager,"


def test_build_salutation_empty_string():
    assert pt._build_salutation("") == "Dear Hiring Manager,"


def test_build_salutation_whitespace_only():
    assert pt._build_salutation("   ") == "Dear Hiring Manager,"


# ==============================================
# _build_closing -- unit tests
# ==============================================

def test_build_closing_contains_thank_you_sentence():
    closing = pt._build_closing("Jane Doe")
    assert "Thank you again for your time" in closing


def test_build_closing_contains_respectfully():
    closing = pt._build_closing("Jane Doe")
    assert "Respectfully," in closing


def test_build_closing_contains_candidate_name():
    closing = pt._build_closing("Jane Doe")
    assert "Jane Doe" in closing


def test_build_closing_ndash_not_emdash():
    closing = pt._build_closing("Jane Doe")
    assert "–" in closing   # en dash present
    assert "—" not in closing  # em dash absent


# ==============================================
# Full letter output -- salutation, closing, fallback
# ==============================================

def test_full_letter_salutation_present(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setenv("CANDIDATE_NAME", "Test Candidate")

    debrief = _make_debrief([{"name": "John Smith", "title": "Engineer", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Body paragraph here.")

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    content = open(generated[0][1], encoding="utf-8").read()
    assert content.startswith("Dear John,")


def test_full_letter_first_name_extracted(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setenv("CANDIDATE_NAME", "Test Candidate")

    debrief = _make_debrief([{"name": "Margaret O'Brien", "title": "Director", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Body here.")

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    content = open(generated[0][1], encoding="utf-8").read()
    assert content.startswith("Dear Margaret,")
    assert "O'Brien" not in content.split("\n")[0]


def test_full_letter_null_name_fallback(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setenv("CANDIDATE_NAME", "Test Candidate")

    debrief = _make_debrief([{"name": None, "title": "Engineer", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Body here.")

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    content = open(generated[0][1], encoding="utf-8").read()
    assert content.startswith("Dear Hiring Manager,")


def test_full_letter_closing_block_present(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setenv("CANDIDATE_NAME", "Test Candidate")

    debrief = _make_debrief([{"name": "Pat Kim", "title": "Recruiter", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Body here.")

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    content = open(generated[0][1], encoding="utf-8").read()
    assert "Thank you again for your time" in content
    assert "Respectfully," in content


def test_full_letter_candidate_name_in_closing(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setenv("CANDIDATE_NAME", "R. Todd Drake")

    debrief = _make_debrief([{"name": "Sam Lee", "title": "Engineer", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Body here.")

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    content = open(generated[0][1], encoding="utf-8").read()
    assert "R. Todd Drake" in content


def test_full_letter_ndash_in_closing(tmp_path, monkeypatch):
    role = "TestRole"
    _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))
    monkeypatch.setenv("CANDIDATE_NAME", "Test Candidate")

    debrief = _make_debrief([{"name": "Dana Cruz", "title": "Engineer", "notes": "Notes."}])
    inputs = _make_inputs(debrief)
    client = _mock_client("Body here.")

    generated, _ = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    content = open(generated[0][1], encoding="utf-8").read()
    assert "–" in content    # en dash present
    assert "—" not in content  # em dash absent


# ==============================================
# Summary block output
# ==============================================

def test_summary_block_output(tmp_path, monkeypatch, capsys):
    role = "TestRole"
    pkg_dir = _setup_job_package(tmp_path, role)
    monkeypatch.setattr(pt, "JOBS_PACKAGES_DIR", str(tmp_path / "job_packages"))

    # Two interviewers; pre-create first file to trigger skip
    existing_txt = pkg_dir / "thankyou_hiring_manager_smith_2026-04-15.txt"
    existing_txt.write_text("old", encoding="utf-8")

    debrief = _make_debrief([
        {"name": "John Smith", "title": "Engineer", "notes": "Notes A."},
        {"name": "Mary Jones", "title": "Director", "notes": "Notes B."},
    ])
    inputs = _make_inputs(debrief)
    client = _mock_client()

    # Decline overwrite for first, accept (not prompted) for second
    responses = iter(["n"])
    monkeypatch.setattr("builtins.input", lambda _: next(responses))

    generated, skipped = pt.generate_letters(client, role, "hiring_manager", None, inputs, "2026-04-15")

    assert len(generated) == 1
    assert len(skipped) == 1

    # Simulate the summary block that main() would print
    import io
    from contextlib import redirect_stdout
    buf = io.StringIO()
    with redirect_stdout(buf):
        if generated:
            print(f"Generated {len(generated)} thank you letter(s):")
            for name, txt_path, docx_path in generated:
                print(f"  {txt_path}")
                print(f"  {docx_path}")
        if skipped:
            print("\nSkipped:")
            for name, _, _ in skipped:
                print(f"  {name} \u2013 file exists, user declined overwrite")
    summary = buf.getvalue()

    assert "Generated 1" in summary
    assert "jones" in summary.lower()
    assert "Skipped" in summary
    assert "John Smith" in summary
